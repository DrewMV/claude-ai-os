"""
md_to_docx.py — Vault-wide Markdown to DOCX converter.

Usage:
  python _scripts/md_to_docx.py <file.md> [file2.md ...]
  python _scripts/md_to_docx.py Workspaces/Work/Projects/nerc-cip/templates/configuration-management-plan.md

Output: <same folder as input>/<OriginalName>.docx
Supports: headings, tables, blockquotes (NERC callout boxes), bullets,
          numbered lists, checklists, bold/italic, inline code, code blocks,
          horizontal rules, YAML frontmatter (stripped).
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        tag = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            tag.set(qn(f"w:{k}"), v)
        tcBorders.append(tag)
    tcPr.append(tcBorders)


# ---------------------------------------------------------------------------
# Inline markup parser
# ---------------------------------------------------------------------------

def apply_inline(text, para, base_size=Pt(11)):
    """Add runs to *para* honouring **bold**, *italic*, `code`."""
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)')
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
            r.font.size = base_size
        elif part.startswith("*") and part.endswith("*"):
            r = para.add_run(part[1:-1])
            r.italic = True
            r.font.size = base_size
        elif part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
        else:
            r = para.add_run(part)
            r.font.size = base_size


# ---------------------------------------------------------------------------
# Style setup
# ---------------------------------------------------------------------------

def setup_styles(doc):
    s = doc.styles

    n = s["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)

    h1 = s["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    h2 = s["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)

    h3 = s["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------

def render_callout(doc, bq_lines):
    """Render blockquote lines as a blue-bordered NERC callout box."""
    text = " ".join(ln.lstrip("> ").strip() for ln in bq_lines if ln.strip())
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, "EBF3FB")
    set_cell_border(cell,
        top={"val": "single", "sz": "4", "color": "2E75B6"},
        bottom={"val": "single", "sz": "4", "color": "2E75B6"},
        left={"val": "single", "sz": "18", "color": "2E75B6"},
        right={"val": "single", "sz": "4", "color": "2E75B6"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    apply_inline(text, p, base_size=Pt(10))
    # Colour the opening "NERC CIP" label dark navy
    if p.runs and "NERC CIP" in p.runs[0].text:
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p.runs[0].bold = True
    doc.add_paragraph()


def render_table(doc, table_lines):
    """Render pipe-table lines as a formatted Word table."""
    rows = []
    for tl in table_lines:
        cols = [c.strip() for c in tl.strip().strip("|").split("|")]
        rows.append(cols)
    # Drop separator row (cells that are all dashes/colons/spaces)
    data = [r for r in rows if not all(re.match(r"^[\s\-:]+$", c) for c in r)]
    if not data:
        return
    col_count = max(len(r) for r in data)
    data = [r + [""] * (col_count - len(r)) for r in data]

    t = doc.add_table(rows=0, cols=col_count)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ri, row_data in enumerate(data):
        row = t.add_row()
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ri == 0:                          # header row
                set_cell_bg(cell, "2E75B6")
                r = p.add_run(cell_text)
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
            else:
                if ri % 2 == 0:
                    set_cell_bg(cell, "F2F7FB")
                apply_inline(cell_text, p, base_size=Pt(10))

    doc.add_paragraph()


def render_code_block(doc, code_lines):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("\n".join(code_lines))
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def render_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def convert(md_path: Path) -> Path:
    out_path = md_path.with_suffix(".docx")

    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    # Strip YAML frontmatter
    i = 0
    if lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---":
            i += 1
        i += 1

    in_code = False
    code_lines: list[str] = []
    bq_lines: list[str] = []
    table_lines: list[str] = []

    def flush_bq():
        if bq_lines:
            render_callout(doc, bq_lines)
        bq_lines.clear()

    def flush_table():
        if table_lines:
            render_table(doc, table_lines)
        table_lines.clear()

    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            flush_bq(); flush_table()
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                render_code_block(doc, code_lines)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            flush_table()
            bq_lines.append(line)
            i += 1
            continue
        elif bq_lines:
            flush_bq()

        # Table
        if line.strip().startswith("|") and "|" in line:
            flush_bq()
            table_lines.append(line)
            i += 1
            continue
        elif table_lines:
            flush_table()

        # Horizontal rule
        if re.match(r"^---+\s*$", line.strip()):
            render_hr(doc)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 3")
            p = doc.add_paragraph(style=style)
            apply_inline(text, p)
            i += 1
            continue

        # Checkbox list
        m = re.match(r"^(\s*)-\s+\[([ xX])\]\s+(.*)", line)
        if m:
            checked = m.group(2).strip().lower() == "x"
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run("☑  " if checked else "☐  ")
            r.font.size = Pt(11)
            apply_inline(m.group(3), p)
            i += 1
            continue

        # Bullet list
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            p.paragraph_format.space_after = Pt(2)
            apply_inline(m.group(2), p)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            p.paragraph_format.space_after = Pt(2)
            apply_inline(m.group(2), p)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(6)
        apply_inline(line, p)
        i += 1

    # Flush anything pending
    flush_bq()
    flush_table()
    if in_code and code_lines:
        render_code_block(doc, code_lines)

    doc.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python _scripts/md_to_docx.py <file.md> [file2.md ...]")
        sys.exit(1)

    vault_root = Path(__file__).parent.parent

    for arg in sys.argv[1:]:
        p = Path(arg)
        if not p.is_absolute():
            p = vault_root / p
        if not p.exists():
            print(f"Not found: {p}")
            continue
        out = convert(p)
        print(f"OK  {out.relative_to(vault_root)}")
