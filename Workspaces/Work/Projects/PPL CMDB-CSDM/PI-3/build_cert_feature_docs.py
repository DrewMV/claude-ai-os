"""
Render the PI-3 certification feature stories-and-dependencies markdown files as
shareable Word (.docx) documents, in the team house style.

Reads the markdown files directly (single source of truth — keep the .md authoritative;
this script only renders). Handles the constructs those files actually use: YAML
frontmatter (→ title page), H1/H2/H3 headings (H3 = one story/dependency per page),
blockquote callouts, pipe tables, bullets, [ ] checkbox bullets, fenced code blocks
(sequence diagrams), and inline **bold** / `code` / [[wikilinks]] / [text](links).

Run:     python build_cert_feature_docs.py
Output:  Certify-Computers-Stories-1517029-2026-07-27.docx
         Certify-Servers-Stories-1517032-2026-07-27.docx
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

EDITION = "2026-07-27"

DARK   = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x6D, 0xB4)
MUTED  = RGBColor(0x6B, 0x72, 0x80)
RED    = RGBColor(0xC0, 0x39, 0x2B)

FILES = [
    ("pi3-cert-1517029-certify-computers-stories-dependencies.md",
     "Certify-Computers-Stories-1517029-" + EDITION + ".docx"),
    ("pi3-cert-1517032-certify-servers-stories-dependencies.md",
     "Certify-Servers-Stories-1517032-" + EDITION + ".docx"),
    ("pi3-cert-1517037-certify-databases-stories-dependencies.md",
     "Certify-Databases-Stories-1517037-" + EDITION + ".docx"),
    ("pi3-cert-1517040-certify-network-devices-stories-dependencies.md",
     "Certify-Network-Devices-Stories-1517040-" + EDITION + ".docx"),
]


# --------------------------------------------------------------------------- PARSE
def split_frontmatter(text):
    """Return (frontmatter dict-ish lines, body) — frontmatter is between leading ---/---."""
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            fm = text[3:end].strip()
            body = text[end + 4:].lstrip("\n")
            meta = {}
            for line in fm.splitlines():
                m = re.match(r"^([a-zA-Z0-9_-]+):\s*(.*)$", line)
                if m and m.group(2):
                    meta[m.group(1)] = m.group(2).strip().strip('"')
            return meta, body
    return {}, text


def normalize_inline(s):
    """Strip wikilink/link syntax down to display text (keeps **bold** and `code` markers)."""
    s = re.sub(r"\[\[([^\]]+)\]\]", lambda m: m.group(1).split("|")[-1], s)   # [[a|b]] -> b
    s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)                             # [t](u) -> t
    return s


def add_inline(paragraph, text, base_color=None):
    """Add text to a paragraph, honoring **bold** and `code` spans."""
    text = normalize_inline(text)
    for tok in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        else:
            r = paragraph.add_run(tok)
        if base_color is not None:
            r.font.color.rgb = base_color


# --------------------------------------------------------------------------- RENDER
def style_base(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(10.5)


def heading(doc, text, size, color=DARK, space_before=8, space_after=4, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    add_inline(p, text)
    for r in p.runs:
        r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, checkbox=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    if checkbox:
        p.add_run("\u2610  ").font.color.rgb = ACCENT
    else:
        p.add_run("\u2022  ").font.color.rgb = ACCENT
    add_inline(p, text)
    return p


def callout(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    add_inline(p, "  ".join(lines))
    for r in p.runs:
        r.italic = True
        if r.font.color.rgb is None:
            r.font.color.rgb = MUTED
    return p


def code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run("\n".join(lines))
    r.font.name = "Consolas"; r.font.size = Pt(8.5); r.font.color.rgb = DARK
    return p


def emit_table(doc, rows):
    cols = max(len(r) for r in rows)
    rows = [r + [""] * (cols - len(r)) for r in rows]
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            add_inline(cell.paragraphs[0], val)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(8.5)
                if i == 0:
                    r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_separator(line):
    return bool(re.match(r"^\|?[\s:\-|]+\|?$", line)) and "-" in line


# --------------------------------------------------------------------------- BUILD
def build(md_path):
    with open(md_path, encoding="utf-8") as f:
        meta, body = split_frontmatter(f.read())

    doc = Document()
    style_base(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)

    # Title page block from frontmatter
    feature = meta.get("feature", "").strip('"')
    title = doc.add_paragraph()
    r = title.add_run("PPL CMDB-CSDM — PI-3 Data Certification"); r.bold = True
    r.font.size = Pt(20); r.font.color.rgb = DARK
    if feature:
        sub = doc.add_paragraph()
        r = sub.add_run("Feature " + feature); r.font.size = Pt(13); r.font.color.rgb = ACCENT
    metaline = doc.add_paragraph()
    bits = ["DRAFT · " + EDITION]
    if meta.get("objective"): bits.append(meta["objective"])
    if meta.get("sprint"):    bits.append("Sprint " + meta["sprint"].strip('"'))
    if meta.get("co6-due"):   bits.append("CO6 due " + meta["co6-due"])
    r = metaline.add_run("  ·  ".join(bits))
    r.font.size = Pt(9); r.font.color.rgb = MUTED; r.italic = True

    # Body — line-by-line state machine
    lines = body.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            code_block(doc, buf)
            i += 1
            continue

        # table (consecutive pipe lines)
        if stripped.startswith("|"):
            buf = []
            while i < n and lines[i].strip().startswith("|"):
                if not is_separator(lines[i]):
                    buf.append(parse_table_row(lines[i]))
                i += 1
            if buf:
                emit_table(doc, buf)
            continue

        # blockquote (consecutive > lines)
        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            callout(doc, [b for b in buf if b])
            continue

        # headings
        if stripped.startswith("### "):
            heading(doc, stripped[4:], 14, DARK, space_before=0, page_break=True)
            i += 1; continue
        if stripped.startswith("## "):
            heading(doc, stripped[3:], 13, ACCENT)
            i += 1; continue
        if stripped.startswith("# "):
            heading(doc, stripped[2:], 16, DARK)
            i += 1; continue

        # bullets
        m = re.match(r"^\s*-\s+\[ \]\s+(.*)$", line)
        if m:
            bullet(doc, m.group(1), checkbox=True); i += 1; continue
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            bullet(doc, m.group(1)); i += 1; continue

        # horizontal rule / blank
        if stripped == "---" or stripped == "":
            i += 1; continue

        # normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_inline(p, stripped)
        i += 1

    return doc


def save_safe(doc, path):
    try:
        doc.save(path); print("Saved", os.path.basename(path))
    except PermissionError:
        base, ext = os.path.splitext(path)
        alt = base + "-new" + ext
        doc.save(alt); print("Target locked — saved to", os.path.basename(alt))


here = os.path.dirname(os.path.abspath(__file__))
for md, out in FILES:
    save_safe(build(os.path.join(here, md)), os.path.join(here, out))
