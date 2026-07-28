"""
CMDB Health configuration — SAFe stories, Word (.docx) output.

Renders the 7 stories from cmdb-health-completeness-correctness-stories.md as a
shareable Word document: title page + parent feature + conventions + story index,
then one story per page (SAFe: User Story + Acceptance Criteria + Notes).

Content mirrors the markdown source (single source of truth) — keep them in sync.

Run:     python build_cmdb_health_stories_doc.py
Output:  CMDB-Health-CI-Stories-2026-07-20.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

EDITION = "2026-07-20"
OUTFILE = "CMDB-Health-CI-Stories-" + EDITION + ".docx"

DARK   = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x6D, 0xB4)
MUTED  = RGBColor(0x6B, 0x72, 0x80)
RED    = RGBColor(0xC0, 0x39, 0x2B)

# --------------------------------------------------------------------------- DATA
PARENT = dict(
    name="[Enabler · Configuration] CMDB Health & Data Quality — KPI Configuration",
    owner="Product Management (Sonika Das) + PO (Joe Dames); driven by Manuel Vazquez (acting PO-support / SM)",
    target="PI-3 (config can start late PI-2 2.5/2.6) — confirm with Sonika",
    priority="Completeness = Medium-High (maps to active audit-readiness remediation); Correctness = Medium (not holdback-gating)",
    benefit=("If Completeness and Correctness are configured per class from the agreed managed-attribute set and scored, "
             "then CMDB population and data-integrity gaps become measurable against target and auditable — supporting "
             "monthly audit-readiness reporting and trust in incident routing / change approval / reporting."),
    scope=("Completeness and Correctness KPIs only. Excludes Compliance (SOX / ESS-02 audits) and value-validation. "
           "Correctness excludes Business Application (not discovery-based)."),
)

CONVENTIONS = [
    "Persona (scoped to these stories): “CMDB Configuration Management Process Owner” = Manuel Vazquez (acting / PO-support), pending confirmation. System-of-record Process Owner & CCB Chairperson remains Josh Sterling.",
    "Target: working 85–90% score per class/KPI — confirm formally with Sonika; confirm whether targets differ per class.",
    "Acceptance lane (non-SOW): PM Sonika Das accepts targets; PO Joe Dames sign-off at Sprint Review; CCB informed (visibility only). Tag NON-SOW / Quality-Governance.",
    "Staleness threshold N (Correctness) and orphan rule (one per class, manual): agreed with the CCB Class Manager — placeholders marked [confirm].",
    "Story vs Enabler: platform configuration → Enabler – Configuration; if the ART doesn't use Enabler type, User Story + Configuration/Quality-Governance tag is fine.",
    "Source: attribute lists, field names, requirements, identification rules and CI counts are from the attached CI Class Reference slides (2026-07-20). All 7 stories must satisfy the team Definition of Done.",
]

# index rows: (id, kpi, class, type, priority, ccb)
INDEX = [
    ("CP-1", "Completeness", "Computer", "Enabler – Configuration", "P2 (confirm)", "Monica Green / Paul Becker"),
    ("CP-2", "Completeness", "Server", "Enabler – Configuration", "P2 (confirm)", "Ray Reuter"),
    ("CP-3", "Completeness", "Database Instance", "Enabler – Configuration", "P2 (confirm)", "Ray Reuter"),
    ("CP-4", "Completeness", "Business Application", "Enabler – Configuration", "P2 (confirm)", "Todd Dierksheide"),
    ("CR-1", "Correctness", "Computer", "Enabler – Configuration", "P3 (confirm)", "Monica Green / Paul Becker"),
    ("CR-2", "Correctness", "Server", "Enabler – Configuration", "P3 (confirm)", "Ray Reuter"),
    ("CR-3", "Correctness", "Database Instance", "Enabler – Configuration", "P3 (confirm)", "Ray Reuter"),
]

STORIES = [
    dict(
        id="CP-1", title="Configure CMDB Health Completeness — Computer (Physical/Virtual)",
        kpi="Completeness", typ="Enabler – Configuration (NON-SOW)", priority="P2 (confirm)",
        table="cmdb_ci_computer", count="19,901 active", ccb="Monica Green / Paul Becker",
        scope="Confirm full-population vs Physical/Virtual split.",
        as_="the CMDB Configuration Management Process Owner",
        want="the CMDB Health Completeness KPI configured from the Computer managed attributes and scored",
        so_that="computer population gaps are measurable against target and feed audit-readiness reporting",
        acs=[
            "Recommended attribute set configured as the completeness basis (no attributes marked Required — acceptable): Assigned To, CI Owner (owned_by), Technical Owner Group (u_technical_owner_group), Support Group, Asset Tag, OS, Location, IP Address, Serial Number.",
            "Completeness scored for cmdb_ci_computer; baseline captured vs target (85–90%, confirm).",
            "Score cross-checks the Computer audit dashboard gaps (Missing CI Owner, Assigned To empty, Location empty, Missing IP) — coordinated with spike 1480114 so remediation isn't double-counted.",
            "Any required relationships for completeness defined (or explicitly none).",
            "PM (Sonika) accepts target; PO (Joe) sign-off; CCB informed; logged in quality register. Team Definition of Done met.",
        ],
        notes="ESS-02-audited attributes (Assigned To, Location, Serial Number) are natural completeness priorities. Location gap has a documented HR-system root cause; offline physical devices limit discovered-attribute population.",
    ),
    dict(
        id="CP-2", title="Configure CMDB Health Completeness — Server (Windows/Linux)",
        kpi="Completeness", typ="Enabler – Configuration (NON-SOW)", priority="P2 (confirm)",
        table="cmdb_ci_server", count="3,917 active", ccb="Ray Reuter",
        scope="",
        as_="the CMDB Configuration Management Process Owner",
        want="the CMDB Health Completeness KPI configured from the Server managed attributes and scored",
        so_that="server population gaps are measurable against target and audit-readiness on the server class is sustained",
        acs=[
            "Required set configured (proposed *, pending CCB): CI Owner (owned_by), Technical Owner Group (u_technical_owner_group), Support Group, Value Stream (u_value_stream — confirm field), Classification (u_classification — confirm), Location.",
            "Recommended set configured: Asset Tag, SOX Type, IP Address.",
            "Completeness scored for cmdb_ci_server; baseline vs target (85–90%, confirm).",
            "Score reflects recent remediation (Location 44%→100%, IP 43%→91%) and coordinates with the Servers audit spike.",
            "Required/Recommended designations and the proposed * flags approved by Ray Reuter.",
            "Any required relationships defined (or explicitly none).",
            "PM (Sonika) accepts target; PO (Joe) sign-off; CCB informed; logged in quality register. Team Definition of Done met.",
        ],
        notes="Server is the most-defined class (6 Required attrs) and already audit-ready on its ESS-02 attributes — this formalizes and monitors that against target.",
    ),
    dict(
        id="CP-3", title="Configure CMDB Health Completeness — Database Instance",
        kpi="Completeness", typ="Enabler – Configuration (NON-SOW)", priority="P2 (confirm)",
        table="cmdb_ci_database", count="3,113 active", ccb="Ray Reuter",
        scope="Child classes (Oracle / MS SQL) inherit.",
        as_="the CMDB Configuration Management Process Owner",
        want="the CMDB Health Completeness KPI configured from the Database managed attributes and scored",
        so_that="database population gaps are measurable against target",
        acs=[
            "Recommended attribute set configured (no attributes marked Required — acceptable), using slide field names: CI Owner (managed_by), Technical Owner Group (managed_by_group), Support Group, Environment (environment — confirm), Location.",
            "Completeness scored for cmdb_ci_database; baseline vs target (85–90%, confirm).",
            "Reconcile the attribute set with the audit dashboard first (BLOCKING) — agree one authoritative Database attribute set with Ray Reuter so Completeness and the audit dashboard don't measure different things (see mismatch table below).",
            "Any required relationships defined (e.g., hosted-on Server) or explicitly none.",
            "PM (Sonika) accepts target; PO (Joe) sign-off; CCB informed; logged in quality register. Team Definition of Done met.",
        ],
        notes="DB uses the managed_by / managed_by_group mapping. Discovered-attribute completeness is bounded by DB discovery coverage (credential risk 1444864; Oracle-KY gap); ownership fields are manual / Data-Cert populated.",
        mismatch_title="CP-3 attribute mismatch — Database audit dashboard (spike 1480113) vs. managed-attribute slide (overlap on only 3 of 8 fields):",
        mismatch=[
            ("Attribute", "On slide?", "Slide Audit", "On dashboard?", "Gap"),
            ("CI Owner (managed_by)", "Yes", "T", "Yes – Missing CI Owner", "10"),
            ("Support Group", "Yes", "F", "Yes – Missing Support Group", "15"),
            ("Environment", "Yes", "F", "Yes – Missing Environment", "1,187"),
            ("Technical Owner Group (managed_by_group)", "Yes", "F", "No – not audited", "—"),
            ("Location", "Yes", "T", "No – not audited", "—"),
            ("Value Stream (Business Unit)", "NOT on slide", "—", "Yes – Missing Value Stream", "3,113 (100%)"),
            ("SOX Type", "NOT on slide", "—", "Yes – SOX Type empty", "75"),
            ("Approval Group", "NOT on slide", "—", "Yes – Missing Approval Group", "3,110 (~99.9%)"),
        ],
    ),
    dict(
        id="CP-4", title="Configure CMDB Health Completeness — Business Application",
        kpi="Completeness", typ="Enabler – Configuration (NON-SOW)", priority="P2 (confirm)",
        table="cmdb_ci_business_app", count="2,148 records", ccb="Todd Dierksheide",
        scope="Business Applications are not discovered — population is manual / Data-Certification sourced.",
        as_="the CMDB Configuration Management Process Owner",
        want="the CMDB Health Completeness KPI configured from the Business Application managed attributes and scored",
        so_that="BA population gaps (the largest audit-readiness gaps) are measurable against target — despite BAs not being discoverable",
        acs=[
            "Recommended attribute set configured (no attributes marked Required except Name/identifier — acceptable), using slide field names: CI Owner (managed_by), Technical Owner Group (managed_by_group), Business Owner (owned_by), Approval Group (u_approval_group), Support Group (support_group), Asset Tag (asset_tag), Value Stream/Business Unit (business_unit), SOX Type (u_sox_type), Data Classification (data_classification).",
            "Completeness scored for cmdb_ci_business_app; baseline vs target (85–90%, confirm) — population is manual / Data-Certification sourced, not Discovery.",
            "Coordinated with the BA audit spike (1480111). Reconcile one gap: the BA audit dashboard also audits Missing Recovery Tier (1,820), which is NOT on the BA managed-attribute slide — confirm whether Recovery Tier joins the completeness set.",
            "Approved by Todd Dierksheide (CCB Class Manager, Business Application).",
            "PM (Sonika) accepts target; PO (Joe) sign-off; CCB informed; logged in quality register. Team Definition of Done met.",
        ],
        notes="BA slide (9 managed attributes) matches the audit dashboard on 9 of 10 audits — only Recovery Tier is audited-but-not-tabled (contrast the larger Database mismatch in CP-3). Respect the BA enhancement pause (6/10): mass BA edits are logged as stories, not done inline. Ties to Data Certification (SOW 1.2), CO5 1.1c, and the BA data-population cluster (1475582/1475584/1475585/1478286/1474892). Completeness counterpart to HL-4 (certification-based lifecycle/staleness).",
    ),
    dict(
        id="CR-1", title="Configure CMDB Health Correctness — Computer (Physical/Virtual)",
        kpi="Correctness", typ="Enabler – Configuration (NON-SOW)", priority="P3 (confirm)",
        table="cmdb_ci_computer", count="19,901 active", ccb="Monica Green / Paul Becker",
        scope="Confirm full-population vs Physical/Virtual split.",
        as_="the CMDB Configuration Management Process Owner",
        want="the CMDB Health Correctness KPI (Duplicate · Orphan · Stale) configured and scored for the Computer class",
        so_that="computer data integrity is measurable against target and remediation is targeted rather than a raw gap count",
        acs=[
            "Duplicate sub-metric active and scored using the PPL Computers identification rule (Name + serial_number, Pri 100); duplicate set reviewed and routed to the de-duplication / CI remediation task. (Serial coverage strong — validate Name uniqueness.)",
            "Orphan rule defined and agreed with CCB Class Managers. Risk: many end-user computers legitimately have no relationships → a naive rule over-flags; record what makes a Computer a non-orphan (or exempt Computer from the orphan sub-metric).",
            "Stale rule set — Computer not updated by SCCM/Discovery within N days ([confirm], e.g. 30/45) is flagged. Risk: offline physical devices skew staleness (SCCM precedence 1348712/716/717; retired-computer 1402790). Stale CIs feed lifecycle (HL-2).",
            "Computer Correctness score visible on the dashboard / baseline (ties 1436581 / 1470837), scored vs target (85–90%, confirm).",
            "PM (Sonika) accepts target; PO (Joe) sign-off; CCB informed; logged in quality register. Team Definition of Done met.",
        ],
        notes="Coordinate with the Computer audit spike (1480114) so correctness config and audit-gap remediation don't double-count.",
    ),
    dict(
        id="CR-2", title="Configure CMDB Health Correctness — Server (Windows/Linux)",
        kpi="Correctness", typ="Enabler – Configuration (NON-SOW)", priority="P3 (confirm)",
        table="cmdb_ci_server", count="3,917 active", ccb="Ray Reuter",
        scope="",
        as_="the CMDB Configuration Management Process Owner",
        want="the CMDB Health Correctness KPI (Duplicate · Orphan · Stale) configured and scored for the Server class",
        so_that="server data integrity is measurable against target and remediation is targeted",
        acs=[
            "Duplicate sub-metric active and scored using the PPL Server identification rule — serial_number (100) → serial lookup (200) → name (300). Note: serial_number is the primary identifier though it is NOT in the Server managed-attribute set — validate serial coverage; low coverage inflates duplicates. Duplicate set reviewed and remediated.",
            "Orphan rule = a Server with no relationship to any Application / Business Service, agreed with Ray Reuter. Score expected to improve as Service Mapping (P1) builds relationships (lead currently vacant — see risks).",
            "Stale rule — Server not seen by Discovery / SG-SCCM within N days ([confirm]) is flagged (reconciliation inherited from the Computer/Hardware hierarchy). Stale CIs feed lifecycle (HL-1).",
            "Server Correctness score visible on the dashboard / baseline, scored vs target (85–90%, confirm). (Distinct from the Server Compliance/Completeness audit-ready milestone.)",
            "PM (Sonika) accepts target; PO (Joe) sign-off; CCB informed; logged in quality register. Team Definition of Done met.",
        ],
        notes="Coordinate with the Servers audit spike. Orphan sub-metric is the one most affected by Service Mapping progress.",
    ),
    dict(
        id="CR-3", title="Configure CMDB Health Correctness — Database Instance",
        kpi="Correctness", typ="Enabler – Configuration (NON-SOW)", priority="P3 (confirm)",
        table="cmdb_ci_database", count="3,113 active", ccb="Ray Reuter",
        scope="Child classes (Oracle / MS SQL) add engine-specific rules.",
        as_="the CMDB Configuration Management Process Owner",
        want="the CMDB Health Correctness KPI (Duplicate · Orphan · Stale) configured and scored for the Database Instance class",
        so_that="database data integrity is measurable against target and remediation is targeted",
        acs=[
            "Duplicate sub-metric active and scored using the Database instance rule — serial_number (100) → Edition + name + TCP port(s) (200). Note: the rule is Dependent — identified in the context of its host Server, so duplicate detection is only as correct as host-server identity; validate Edition/name/TCP port populated (identifier attributes, not the governance managed attributes). Duplicate set reviewed and remediated.",
            "Orphan rule = a Database Instance not hosted on a Server (Runs on::Runs absent), agreed with Ray Reuter. A hostless DB is both an identity and orphan problem; expect orphans where Discovery can't authenticate to the host (Oracle-KY gap, RAID).",
            "Stale rule — Database not seen by ServiceNow application probes within N days ([confirm]) is flagged. Dependency: only as good as DB discovery coverage (credential risk 1444864; enhanced DB discovery D2.3; Oracle-KY gap). Stale CIs feed lifecycle (HL-3).",
            "Database Correctness score visible on the dashboard / baseline, scored vs target (85–90%, confirm).",
            "PM (Sonika) accepts target; PO (Joe) sign-off; CCB informed; logged in quality register. Team Definition of Done met.",
        ],
        notes="Coordinate with the Database audit spike. Confirm whether Correctness is scored at the cmdb_ci_database parent or per child class (Oracle / MS SQL).",
    ),
]

TO_CONFIRM = [
    "Intent — CR stories build the ServiceNow Correctness KPI (Duplicate/Orphan/Stale), not value-validation (Compliance). (Joe / Sonika)",
    "Target — 85–90% per class/KPI, or different; recommended-attribute weighting in the deployed release. (Sonika)",
    "Database attribute set (blocking, CP-3) — reconcile the 5 slide attributes vs the 6 audit-dashboard audits (dashboard adds Value Stream / SOX Type / Approval Group; slide adds Technical Owner Group / Location). (Ray Reuter)",
    "BA Recovery Tier (CP-4) — does it join the BA completeness set? (Todd Dierksheide)",
    "Field names — u_value_stream, u_classification (Server), environment (Database), and the managed_by / managed_by_group mapping (Database, BA).",
    "Staleness N per class + orphan rule definitions (Computer over-flag risk); Service Mapping lead reassignment (orphan dependency).",
    "Computer scope — full population vs Physical/Virtual split.",
    "Parent feature & Enabler type — single CMDB Health & Data Quality parent; Enabler type vs User Story + tag.",
    "Reconcile with HL-1/2/3/4 so their bundled KPI ACs point to these stories rather than duplicating.",
]


# --------------------------------------------------------------------------- HELPERS
def style_base(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(10.5)


def label_value(doc, label, value):
    p = doc.add_paragraph()
    r = p.add_run(label + ": "); r.bold = True; r.font.color.rgb = DARK
    p.add_run(value)
    p.paragraph_format.space_after = Pt(2)
    return p


def heading(doc, text, size, color=DARK, space_before=6, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, checkbox=False):
    p = doc.add_paragraph(style="List Bullet")
    if checkbox:
        p.paragraph_format.left_indent = Inches(0.25)
        p.style = doc.styles["Normal"]
        p.add_run("☐  ").font.color.rgb = ACCENT
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def grid_table(doc, rows, header=True, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = t.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = ""
            para = cells[j].paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(8.5)
            if header and i == 0:
                run.bold = True
            if widths and j < len(widths):
                cells[j].width = widths[j]
    return t


# --------------------------------------------------------------------------- BUILD
def build():
    doc = Document()
    style_base(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)

    # Title
    title = doc.add_paragraph()
    r = title.add_run("PPL CMDB-CSDM — CMDB Health Configuration"); r.bold = True
    r.font.size = Pt(20); r.font.color.rgb = DARK
    sub = doc.add_paragraph()
    r = sub.add_run("SAFe Stories · Completeness & Correctness KPIs")
    r.font.size = Pt(13); r.font.color.rgb = ACCENT
    meta = doc.add_paragraph()
    r = meta.add_run("DRAFT · " + EDITION + " · NON-SOW Quality Governance · Source: CI Class Reference slides (2026-07-20)")
    r.font.size = Pt(9); r.font.color.rgb = MUTED; r.italic = True

    doc.add_paragraph(
        "7 SAFe-compliant stories to configure and score two ServiceNow CMDB Health KPIs: 4 Completeness "
        "(Computer, Server, Database, Business Application) + 3 Correctness (Computer, Server, Database). "
        "Business Application is Completeness-only — it is not a discovered CI, so the Correctness sub-metrics "
        "(Duplicate / Orphan / Stale) do not apply; its freshness is certification-based (HL-4)."
    )

    # Parent feature
    heading(doc, "Parent Feature (draft)", 13)
    label_value(doc, "Feature", PARENT["name"])
    label_value(doc, "Owner / approver", PARENT["owner"])
    label_value(doc, "Target iterations", PARENT["target"])
    label_value(doc, "WSJF / priority", PARENT["priority"])
    label_value(doc, "Benefit hypothesis", PARENT["benefit"])
    label_value(doc, "Scope / non-goals", PARENT["scope"])

    # Conventions
    heading(doc, "Conventions (apply to all 7 stories)", 13)
    for c in CONVENTIONS:
        bullet(doc, c)

    # Index
    heading(doc, "Story index", 13)
    rows = [("ID", "KPI", "CI Class", "Type", "Priority", "CCB Class Mgr")] + INDEX
    grid_table(doc, rows)

    # Stories, one per page
    for s in STORIES:
        doc.add_page_break()
        heading(doc, s["id"] + " — " + s["title"], 15, DARK, space_before=0)
        meta = doc.add_paragraph()
        rr = meta.add_run("%s  ·  Priority: %s  ·  KPI: %s" % (s["typ"], s["priority"], s["kpi"]))
        rr.font.size = Pt(9); rr.font.color.rgb = MUTED
        label_value(doc, "Scope", "%s (%s)%s" % (
            s["table"], s["count"], ("  —  " + s["scope"]) if s["scope"] else ""))
        label_value(doc, "CCB Class Manager", s["ccb"])
        label_value(doc, "ADO ID / State", "(fill on creation) / Draft")

        heading(doc, "User Story", 12, ACCENT)
        p = doc.add_paragraph()
        p.add_run("As ").bold = False
        p.add_run(s["as_"]).bold = True
        p.add_run(",\nI want ")
        p.add_run(s["want"]).bold = True
        p.add_run(",\nso that ")
        p.add_run(s["so_that"]).bold = True
        p.add_run(".")

        heading(doc, "Acceptance Criteria", 12, ACCENT)
        for ac in s["acs"]:
            bullet(doc, ac, checkbox=True)

        if s.get("mismatch"):
            heading(doc, s["mismatch_title"], 10, RED, space_before=6, space_after=3)
            grid_table(doc, s["mismatch"])

        heading(doc, "Notes", 12, ACCENT)
        doc.add_paragraph(s["notes"])

    # Closing: To confirm
    doc.add_page_break()
    heading(doc, "To confirm before creating in ADO", 14, DARK, space_before=0)
    for t in TO_CONFIRM:
        bullet(doc, t)

    return doc


def save_safe(doc, path):
    try:
        doc.save(path); print("Saved", os.path.basename(path))
    except PermissionError:
        base, ext = os.path.splitext(path)
        alt = base + "-new" + ext
        doc.save(alt); print("Target locked — saved to", os.path.basename(alt))


here = os.path.dirname(os.path.abspath(__file__))
save_safe(build(), os.path.join(here, OUTFILE))
