"""
Build the CMDB Data Dictionary (CMP Stage 1) — CCB Decision Sheet as a Word deliverable.

Emits: Data-Dictionary-CCB-Decision-Sheet.docx  (landscape, matches CMP Stage 1 styling)
Mirrors the content of data-dictionary-ccb-decision-sheet.md (vault-native copy).

Run:  python build_decision_sheet.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

HEADER_FILL = "1F3864"   # dark blue (matches CMP Stage 1)
WARN_FILL = "FCE4D6"     # light orange band
RED = RGBColor(0xC0, 0x00, 0x00)

COLS = ["#", "Attribute", "Decision", "Context", "Team suggestion (advisory)", "CCB ruling"]
COL_W = [0.4, 1.8, 1.0, 3.0, 2.5, 1.5]  # inches (landscape usable ~10.2")

# (class name, class manager, servicenow table, [ (num, attribute, decision, context, suggestion) ])
CLASSES = [
    ("Business Application", "Todd Dierksheide", "cmdb_ci_business_app", [
        (1, "Application Category", "Mandatory?", "Choice; auth = EA/LeanIX", "Mandatory = Y (core classification)"),
        (2, "Life Cycle Stage", "Mandatory?", "Choice (CSDM stages); auth = EA", "Mandatory = Y (lifecycle reporting)"),
        (3, "IT Application Owner", "Mandatory?", "sys_user; via Data Certification", "Mandatory = Y (ownership / audit)"),
        (4, "Business Owner", "Mandatory?", "sys_user; via Data Certification", "Mandatory = Y (Dec audit remediation)"),
    ]),
    ("Service Instance", "*** NONE ASSIGNED ***", "cmdb_ci_service_auto", [
        (5, "Name", "Owner?", "no dedicated CCB Class Mgr mapped", "Assign class owner (EA / Platform)"),
        (6, "Environment", "Owner?", "Prod / QA / Dev / Test", "Same owner as #5"),
        (7, "Service Owner (owned_by)", "Mandatory?", "sys_user", "Mandatory = Y"),
        (8, "ITSM Support Group", "Mandatory?", "sys_user_group", "Mandatory = Y (governance)"),
        (9, "Change Approval Group", "Mandatory?", "sys_user_group", "Mandatory = Y (governance)"),
    ]),
    ("Computer", "Monica Green (Physical) / Paul Becker (Virtual)", "cmdb_ci_computer", [
        (10, "Assigned To", "Mandatory?", "sys_user; Asset Mgmt", "PPL call"),
        (11, "Operating System", "Audited?", "SCCM/Intune authoritative", "Align to Dec audit scope"),
        (12, "MAC Address", "Audited?", "Endpoint connector / Discovery", "Align to Dec audit scope"),
    ]),
    ("Server (parent)", "Ray Reuter", "cmdb_ci_server", [
        (13, "Operational Status", "Values?", "Build / Operational / Non-Operational / Retired", "Confirm full list vs OOB choices"),
        (14, "IP Address", "Audited?", "Discovery authoritative", "Align to Dec audit scope"),
        (15, "MAC Address", "Audited?", "Discovery authoritative", "Align to Dec audit scope"),
        (16, "CPU Count", "Audited?", "vCenter override (else Discovery)", "Align to Dec audit scope"),
        (17, "RAM (MB)", "Audited?", "vCenter override (else Discovery)", "Align to Dec audit scope"),
    ]),
    ("Windows Server (child of Server)", "Ray Reuter", "cmdb_ci_win_server", [
        (18, "OS Version", "Audited?", "SCCM overrides Discovery", "Align to Dec audit scope"),
        (19, "OS Domain", "Audited?", "SCCM overrides Discovery", "Align to Dec audit scope"),
        (20, "OS Service Pack", "Audited?", "SCCM overrides Discovery", "Align to Dec audit scope"),
        (21, "ITSM Support Group", "Mandatory?", "Wintel Infrastructure Ops; not auto-writable", "Mandatory = Y (governance)"),
        (22, "Change Approval Group", "Mandatory?", "Wintel CAB; not auto-writable", "Mandatory = Y (governance)"),
    ]),
    ("Linux Server (child of Server)", "Ray Reuter", "cmdb_ci_linux_server", [
        (23, "OS Architecture", "Audited?", "Discovery via SSH authoritative", "Align to Dec audit scope"),
        (24, "Kernel Release", "Audited?", "Discovery via SSH authoritative", "Align to Dec audit scope"),
        (25, "ITSM Support Group", "Mandatory?", "Unix/Linux Administration Team", "Mandatory = Y (governance)"),
        (26, "Change Approval Group", "Mandatory?", "Linux CAB", "Mandatory = Y (governance)"),
    ]),
    ("Database (parent)", "Ray Reuter", "cmdb_ci_database", [
        (27, "Version", "Audited?", "Discovery authoritative", "Align to Dec audit scope"),
        (28, "TCP Port", "Audited?", "Discovery authoritative", "Align to Dec audit scope"),
        (29, "Operational Status", "Values?", "Operational, Retired, ...", "Confirm full list"),
    ]),
    ("Oracle Database (child of Database)", "Ray Reuter", "cmdb_ci_db_ora_instance", [
        (30, "SID", "Audited?", "Oracle probes", "Align to Dec audit scope"),
        (31, "Oracle Edition", "Audited?", "Oracle probes", "Align to Dec audit scope"),
        (32, "Oracle Version", "Audited?", "Oracle probes", "Align to Dec audit scope"),
        (33, "Support Group", "Mandatory?", "Oracle DBA Team", "Mandatory = Y (governance)"),
    ]),
    ("SQL Server Database (child of Database)", "Ray Reuter", "cmdb_ci_db_mssql_instance", [
        (34, "Instance Name", "Audited?", "Discovery / SCCM", "Align to Dec audit scope"),
        (35, "Clustered", "Audited?", "Discovery authoritative", "Align to Dec audit scope"),
        (36, "SQL Edition", "Audited?", "SCCM authoritative (edition/license)", "Align to Dec audit scope"),
        (37, "TCP Port", "Audited?", "Discovery overrides SCCM", "Align to Dec audit scope"),
        (38, "Support Group", "Mandatory?", "SQL DBA Team", "Mandatory = Y (governance)"),
    ]),
]


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=9, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(0.5))

    doc.add_heading("CMDB Data Dictionary (CMP Stage 1)", level=0)
    doc.add_heading("CCB Decision Sheet", level=1)

    p = doc.add_paragraph()
    p.add_run("Companion to: ").bold = True
    p.add_run("Configuration Management Plan — Stage 1 (Class & Attribute Data Definitions).    ")
    p.add_run("For: ").bold = True
    p.add_run("CMDB CCB vote, 7/21/2026.    ")
    p.add_run("Source draft: ").bold = True
    p.add_run("2026-06-29.")

    doc.add_paragraph(
        "Every cell marked [CCB confirm] in the Data Dictionary is a PPL governance decision left intentionally open. "
        "This sheet lists all 38 of them, grouped by class, so the board can rule on them efficiently rather than "
        "hunting through the document."
    )
    note = doc.add_paragraph()
    note.add_run("How to read this. ").bold = True
    note.add_run(
        "The Team suggestion column is the delivery team's proposed default to speed the vote — it is NOT a decision. "
        "The CCB ruling column is authoritative; the chair captures the outcome there on 7/21."
    )

    leg = doc.add_paragraph()
    leg.add_run("Decision types:").bold = True
    for t in [
        "Mandatory? — confirm whether the attribute is required (Y/N)",
        "Audited? — confirm inclusion in audit / certification scope (Y/N)",
        "Owner? — assign a Certification Owner / Class Manager (none mapped)",
        "Values? — confirm the allowed-value (choice) list",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    # At a glance
    doc.add_heading("At a glance", level=2)
    summary = [
        ("Business Application", "Todd Dierksheide", "4"),
        ("Service Instance", "NONE assigned", "5"),
        ("Computer", "Monica Green (Phys) / Paul Becker (Virt)", "3"),
        ("Server (parent)", "Ray Reuter", "5"),
        ("Windows Server (child)", "Ray Reuter", "5"),
        ("Linux Server (child)", "Ray Reuter", "4"),
        ("Database (parent)", "Ray Reuter", "3"),
        ("Oracle DB (child)", "Ray Reuter", "4"),
        ("SQL Server DB (child)", "Ray Reuter", "5"),
        ("TOTAL", "", "38"),
    ]
    st = doc.add_table(rows=1, cols=3)
    st.style = "Table Grid"
    for j, h in enumerate(["Class", "Class Manager", "# Decisions"]):
        set_cell_text(st.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(st.rows[0].cells[j], HEADER_FILL)
    for name, mgr, n in summary:
        cells = st.add_row().cells
        set_cell_text(cells[0], name, bold=(name == "TOTAL"))
        set_cell_text(cells[1], mgr, color=(RED if mgr == "NONE assigned" else None),
                      bold=(mgr == "NONE assigned"))
        set_cell_text(cells[2], n, bold=(name == "TOTAL"))
    doc.add_paragraph("By type: Mandatory? 14  ·  Audited? 20  ·  Owner? 2  ·  Values? 2.")

    # Warnings
    w = doc.add_paragraph()
    w.add_run("Two things to resolve before the vote:").bold = True
    w1 = doc.add_paragraph(style="List Number")
    w1.add_run("Service Instance has no Class Manager (items 5-6). ").bold = True
    w1.add_run("The board must assign a Certification Owner (EA/Platform) or defer the class; its items cannot be "
               "ratified without one.")
    w2 = doc.add_paragraph(style="List Number")
    w2.add_run("Ray Reuter owns 26 of 38 decisions ").bold = True
    w2.add_run("(all Server + Database classes). Pre-brief him so those classes are not stranded.")

    # Per-class decision tables
    n_class = 0
    for name, mgr, table, items in CLASSES:
        n_class += 1
        doc.add_heading(f"{n_class}. {name}", level=2)
        sub = doc.add_paragraph()
        sub.add_run("Class Manager: ").bold = True
        if mgr.startswith("***"):
            sub.add_run("NO CLASS MANAGER ASSIGNED").font.color.rgb = RED
        else:
            sub.add_run(mgr)
        sub.add_run(f"      Table: {table}").italic = True

        t = doc.add_table(rows=1, cols=len(COLS))
        t.style = "Table Grid"
        for j, col in enumerate(COLS):
            set_cell_text(t.rows[0].cells[j], col, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
            shade_cell(t.rows[0].cells[j], HEADER_FILL)
            t.rows[0].cells[j].width = Inches(COL_W[j])
        for (num, attr, decision, context, suggestion) in items:
            cells = t.add_row().cells
            vals = [num, attr, decision, context, suggestion, ""]
            for j, v in enumerate(vals):
                set_cell_text(cells[j], v, size=9)
                cells[j].width = Inches(COL_W[j])

    # Vote capture
    doc.add_heading("Vote capture", level=2)
    doc.add_paragraph(
        "Record each ruling as Accept (as suggested), Amend (note the change), or Defer (note owner + date). "
        "A class is accepted once all its items are ruled. Roll per-class results into the overall Data Dictionary "
        "acceptance in the CCB minutes."
    )

    out = os.path.join(HERE, "Data-Dictionary-CCB-Decision-Sheet.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    out = build()
    print("Wrote:", out)
    print(f"Classes: {len(CLASSES)} | Items: {sum(len(c[3]) for c in CLASSES)}")
