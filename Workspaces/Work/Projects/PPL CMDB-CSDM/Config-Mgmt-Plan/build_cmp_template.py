"""
Build the reusable Configuration Management Plan TEMPLATE as a Word deliverable.

Emits: Configuration-Management-Plan-Template.docx  (portrait, auto-updating Word TOC)

Mirrors the vault-native _cmp-template.md. Structure validated against official
ServiceNow sources (CSDM Solution Brief, CMDB Data Manager, CMDB Health docs).
Class-centric (Option A): per-class detail lives in section 8; sections 9-11 are
cross-class policy/rules only.

Run:  python build_cmp_template.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

# Frontmatter fields to place at the top of a real CMP doc.
FRONTMATTER = [
    ("Title", "Configuration Management Plan"),
    ("Owner", "<Configuration Management Process Owner>"),
    ("Approver", "<CCB Chairperson>"),
    ("Version", "0.1 (DRAFT)"),
    ("Status", "draft"),
    ("Updated", "<YYYY-MM-DD>"),
]

# (Heading, [body paragraphs]). Body items prefixed "* " render as bullets.
SECTIONS = [
    ("1. Document Control",
     ["Version history (version, date, author, summary of change), approvers / sign-off "
      "(CCB Chair, Process Owner, Architecture), distribution list, and review cycle "
      "(e.g., reviewed quarterly by the CCB)."]),
    ("2. Introduction & Purpose",
     ["Why the plan exists and what “good” looks like for the CMDB. Relationship to ITIL "
      "Service Configuration Management and to the CSDM framework. Intended audience (CCB, "
      "CI Class Owners, Discovery/integration teams)."]),
    ("3. Scope",
     ["* In scope — CI classes governed by this plan: <list>",
      "* Out of scope — <list>",
      "* CSDM 5 domains in scope — Foundation · Ideation & Strategy · Design & Planning · Build & Integration · Service Delivery · Service Consumption · Manage Portfolio",
      "* Environments / business units / data sources covered"]),
    ("4. Objectives & Success Criteria",
     ["Measurable objectives, each tied to a verifiable KPI in section 12/16. "
      "Example: audit coverage ≥ 90% [CCB confirm]."]),
    ("5. References & Related Documents",
     ["CSDM white paper <confirm version>, ServiceNow CMDB docs, contract deliverables, "
      "related governance policies, and the data-dictionary doc(s)."]),
    ("6. Roles & Responsibilities",
     ["* Configuration Management Process Owner — overall CMDB health & policy",
      "* CMDB Manager / CI Data Manager — day-to-day data quality, certification",
      "* CI Class Owners / Class Managers — accountable per class <roster>",
      "* Data Stewards — attribute-level stewardship",
      "* Discovery / Integration owners — authoritative source feeds",
      "* CCB (Change Control Board) — ratifies definitions, policy, precedence rules",
      "* RACI matrix — roles × key activities"]),
    ("7. CMDB / CSDM Architecture",
     ["The seven CSDM 5 domains and what lives in each. CI class hierarchy in scope "
      "(parent/child). How classes map to CSDM 5 domains. Service types "
      "(Business / Technical / Application service)."]),
    ("8. Data Definitions (Class & Attribute Dictionary) — CLASS-CENTRIC",
     ["Authoritative home for all per-class detail. Sections 9-11 state rules; this section "
      "states the values. Link to the data-dictionary doc — do not duplicate. Each class is a "
      "self-contained unit:",
      "* 1. Class Identity & Scope",
      "* 2. Attribute-Level Data Dictionary (incl. Source of Truth, Mandatory, Audited, Certification Owner)",
      "* 3. Data Population & Identification (this class's IRE identifiers + precedence)",
      "* 4. Relationship Architecture (this class's CSDM mappings)"]),
    ("9. Data Sourcing & Reconciliation — Policy (RULES ONLY)",
     ["The IRE philosophy, connector / discovery inventory, and how precedence conflicts are "
      "adjudicated in general. Never restate per-class source-of-truth values — those live in section 8."]),
    ("10. CI Lifecycle Management — Policy (RULES ONLY)",
     ["What each lifecycle stage / operational status means, transition triggers, and staleness, "
      "retirement, and archival rules (ref: CMDB Data Manager). No per-class value lists."]),
    ("11. Relationship & Dependency Standard (RULES ONLY)",
     ["Approved relationship types, top-down vs tag-based vs manual Service Mapping, and mapping "
      "ownership and refresh cadence."]),
    ("12. Data Quality & CMDB Health",
     ["The three metrics — Completeness, Compliance, Correctness. Targets/thresholds per class "
      "[CCB confirm]. Remediation workflow for failing CIs."]),
    ("13. Data Certification & Attestation",
     ["Certification cadence, intake, ownership, attestation flow. OOTB vs custom decision "
      "[CCB confirm]. Evidence retention for audit."]),
    ("14. Governance & Change Control (CCB Operating Model)",
     ["CCB charter, membership, chairperson, quorum. Meeting cadence and decision process. "
      "Change control for data definitions, classes, and precedence rules. Exception / waiver process."]),
    ("15. Security & Compliance Alignment",
     ["SOX scoping, regulatory drivers (e.g., NERC-CIP), security standards. Audit inclusion rules "
      "per class/attribute [CCB confirm]. Access control & data classification."]),
    ("16. Metrics, Reporting & Dashboards",
     ["KPI catalog (health scores, certification %, audit coverage). Dashboards. Report cadence "
      "and audience per report."]),
    ("17. Maintenance & Continual Improvement",
     ["Review/update cadence for this plan. How findings feed back into objectives and the CCB backlog."]),
    ("18. Glossary",
     ["CI, CSDM, IRE, CCB, SGC, Service Mapping, and other key terms."]),
    ("19. Appendices",
     ["* A — RACI matrix",
      "* B — CI class → CSDM domain map",
      "* C — Authoritative-source matrix",
      "* D — KPI definitions",
      "* E — Change / version log detail"]),
]


def add_toc_field(doc):
    """Insert a real, auto-updating Word Table of Contents field (levels 1-2)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose “Update Field” to populate the Table of Contents."
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldSep, placeholder, fldEnd):
        run._r.append(el)


def render_docx():
    doc = Document()
    sec = doc.sections[0]
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(1.0))

    doc.add_heading("Configuration Management Plan", level=0)
    sub = doc.add_paragraph()
    sub.add_run("Reusable Template").italic = True

    intro = doc.add_paragraph()
    intro.add_run("What this is: ").bold = True
    intro.add_run(
        "a reusable skeleton and table of contents for a ServiceNow Configuration Management Plan (CMP), "
        "aligned to the CSDM 5 framework and CMDB governance best practices. Structure validated against "
        "official ServiceNow sources (CSDM 5 white paper, CMDB Data Manager, CMDB Health docs)."
    )
    pat = doc.add_paragraph()
    pat.add_run("Pattern: ").bold = True
    pat.add_run(
        "class-centric (Option A) — all per-class detail lives in section 8; sections 9-11 hold cross-class "
        "policy/rules only and must never restate per-class values."
    )
    conv = doc.add_paragraph()
    conv.add_run("Conventions: ").bold = True
    conv.add_run("[CCB confirm] marks a PPL governance decision requiring CCB ratification. "
                 "<angle brackets> mark fill-ins.")

    # Frontmatter fields table
    doc.add_heading("Document Control Fields", level=1)
    ft = doc.add_table(rows=len(FRONTMATTER), cols=2)
    ft.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(FRONTMATTER):
        c0, c1 = ft.rows[i].cells
        c0.text = ""; c0.paragraphs[0].add_run(k).bold = True
        c1.text = v
        c0.width = Inches(1.8); c1.width = Inches(4.7)

    # Table of Contents (real Word field)
    doc.add_heading("Table of Contents", level=1)
    add_toc_field(doc)

    doc.add_page_break()

    # Body sections
    for heading, body in SECTIONS:
        doc.add_heading(heading, level=1)
        for line in body:
            if line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

    out = os.path.join(HERE, "Configuration-Management-Plan-Template.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    dx = render_docx()
    print("Wrote:", dx)
    print(f"Sections: {len(SECTIONS)} | Frontmatter fields: {len(FRONTMATTER)}")
