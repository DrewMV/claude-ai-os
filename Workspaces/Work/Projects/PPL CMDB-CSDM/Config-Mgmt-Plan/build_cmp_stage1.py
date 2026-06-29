"""
Build the CMDB Configuration Management Plan — Stage 1 (Class & Attribute Data Definitions).

Single structured source -> emits BOTH:
  - configuration-management-plan-stage1.md   (vault-native, version-controlled)
  - Configuration-Management-Plan-Stage1.docx  (CCB review deliverable, landscape)

Run:  python build_cmp_stage1.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DRAFT = "2026-06-24 (draft)"
CONFIRM = "[CCB confirm]"

# Extended data-dictionary columns (order matters for both renderers)
COLS = [
    "Attribute (Display Name)", "Technical Name", "Data Type", "Mandatory",
    "Allowed Values / Reference", "Source of Truth", "Default Value",
    "Audited", "Certification Owner", "Last Reviewed", "Description & Governance",
]

# Each attribute tuple matches COLS order EXCEPT "Last Reviewed" is injected globally as DRAFT.
# tuple = (display, technical, dtype, mandatory, allowed, sot, default, audited, cert_owner, desc)

CLASSES = [
    {
        "name": "Business Application", "kind": "CLASS",
        "display": "Business Application", "table": "cmdb_ci_business_app",
        "domain": "Design & Planning",
        "in_scope": ("The strategic software model — the abstract concept of the application used for tracking "
                     "costs, lifecycle, and business value (e.g., \"ServiceNow HRSD\"). Represents *what* "
                     "functionality is provided."),
        "out_scope": ("Operational, deployed environments (these belong to Service Instances). Underlying physical "
                      "servers or databases."),
        "attrs": [
            ("Name", "name", "String (255)", "Y", "Free text; must be unique", "EA integration (LeanIX) / Manual", "—", "Y",
             "Todd Dierksheide (Class Mgr – Business Application)", "Primary CI identifier; IRE match key."),
            ("Application Category", "category", "Choice", CONFIRM, "HR, Finance, IT, … (choice list)", "EA integration (authoritative)", "—", "Y",
             "Todd Dierksheide (Class Mgr – Business Application)", "High-level grouping. EA integration is authoritative for category."),
            ("Life Cycle Stage", "life_cycle_stage", "Choice", CONFIRM, "CSDM lifecycle stages", "EA integration (authoritative)", "—", "Y",
             "Todd Dierksheide (Class Mgr – Business Application)", "Lifecycle metric sourced from Enterprise Architecture."),
            ("IT Application Owner", "it_application_owner", "Reference (sys_user)", CONFIRM, "Active sys_user record", "Manual / Data Certification", "—", "Y",
             "CI Data Manager (Data Certification)", "Strategic IT leader responsible for the platform lifecycle. Manual updates via Data Certification workflow."),
            ("Business Owner", "business_owner", "Reference (sys_user)", CONFIRM, "Active sys_user record", "Manual / Data Certification", "—", "Y",
             "CI Data Manager (Data Certification)", "Business stakeholder funding/sponsoring the application."),
        ],
        "sources": "Manual entry via Catalog Request, or integration with Enterprise Architecture tools (e.g., LeanIX).",
        "ire": "Name.",
        "precedence": ("Enterprise Architecture integration (if present) is authoritative for Category and Lifecycle "
                       "metrics. Manual updates are permitted for Ownership fields via Data Certification workflows."),
        "upstream": ["Provided by → Business Capability"],
        "downstream": ["Consumes :: Consumed By → Service Instance"],
    },
    {
        "name": "Service Instance", "kind": "CLASS",
        "display": "Service Instance / Application Service",
        "table": "cmdb_ci_service_auto (and subclasses, e.g., cmdb_ci_service_discovered)",
        "domain": "Service Delivery",
        "in_scope": ("The operational, deployed environment — the actual running stack (e.g., \"ServiceNow HRSD - "
                     "PROD\"). Represents *where* the functionality is running."),
        "out_scope": "The abstract software portfolio (Business Application) or individual hardware components.",
        "attrs": [
            ("Name", "name", "String (255)", "Y", "Free text; must be unique", "Service Mapping / Manual", "—", "Y",
             CONFIRM + " — no dedicated CCB Class Manager mapped", "Primary CI identifier."),
            ("Environment", "environment", "Choice", "Y", "Prod, QA, Dev, Test", "Service Mapping / Manual", "—", "Y",
             CONFIRM + " — no dedicated CCB Class Manager mapped", "Deployment environment of the running stack."),
            ("Service Owner", "owned_by", "Reference (sys_user)", CONFIRM, "Active sys_user record", "Manual / Data Manager", "—", "Y",
             "CI Data Manager", "Operational manager responsible for uptime."),
            ("ITSM Support Group", "support_group", "Reference (sys_user_group)", CONFIRM, "Active sys_user_group", "Manual / Data Manager", "—", "Y",
             "CI Data Manager", "Team that handles incidents for this environment."),
            ("Change Approval Group", "change_control", "Reference (sys_user_group)", CONFIRM, "Active sys_user_group", "Manual / Data Manager", "—", "Y",
             "CI Data Manager", "CAB for this environment."),
        ],
        "sources": "ServiceNow Service Mapping (Top-Down), Tag-Based Mapping, or Manual mapping via APIs.",
        "ire": "Inherited / Service Mapping identity (no standalone IRE defined at Stage 1).",
        "precedence": ("ServiceNow Service Mapping is authoritative for dependency structure. ITSM User / Data Manager "
                       "is authoritative for Governance attributes (Support Group, Change Control, Service Owner)."),
        "upstream": ["Consumed by :: Consumes → Business Application", "Used by :: Depends on → Business Service Offering"],
        "downstream": ["Depends on :: Used by → Infrastructure CIs (Servers, Databases)",
                       "Contained by :: Contains → Technology Management Offering"],
    },
    {
        "name": "Computer", "kind": "CLASS",
        "display": "Computer", "table": "cmdb_ci_computer",
        "domain": "Manage Technical Services",
        "in_scope": "End-user computing devices such as laptops, desktops, and workstations.",
        "out_scope": "Servers, Network Gear, Mobile Devices, or IoT devices.",
        "attrs": [
            ("Name", "name", "String (255)", "Y", "Free text", "SCCM/Intune SGC / ACC", "—", "Y",
             "Monica Green (Physical) / Paul Becker (Virtual)", "CI identifier; IRE priority 3."),
            ("Assigned To", "assigned_to", "Reference (sys_user)", CONFIRM, "Active sys_user record", "Asset Mgmt / Manual", "—", "Y",
             "CI Data Manager", "End-user to whom the device is deployed."),
            ("Operating System", "os", "String", "N", "e.g., Windows 11, macOS", "SCCM/Intune SGC (authoritative)", "—", CONFIRM,
             "Monica Green (Physical) / Paul Becker (Virtual)", "OS sourced from endpoint management connector."),
            ("MAC Address", "mac_address", "String", "N", "Valid MAC format", "Endpoint connector / Discovery", "—", CONFIRM,
             "Monica Green (Physical) / Paul Becker (Virtual)", "Network identifier; IRE priority 2."),
            ("Serial Number", "serial_number", "String", "Y", "Hardware serial", "Asset Mgmt / SCCM (authoritative)", "—", "Y",
             "Monica Green (Physical) / Paul Becker (Virtual)", "Hardware serial for asset tracking; IRE priority 1."),
        ],
        "sources": ("Endpoint Management Connectors (Microsoft SCCM/Intune Service Graph Connector, Jamf), "
                    "ServiceNow Agent Client Collector."),
        "ire": "Priority 1: Serial Number. Priority 2: MAC Address. Priority 3: Name.",
        "precedence": ("SCCM/Intune Service Graph Connector is authoritative for OS, OS Version, and Hardware "
                       "specifications. Asset Management / Procurement integrations override for Financial fields "
                       "(Cost Center, PO Number)."),
        "upstream": [],
        "downstream": [],
        "rel_note": ("Typically lacks standard CSDM dependency relationships; usually tied directly to User / Asset "
                     "records."),
    },
    {
        "name": "Server", "kind": "PARENT CLASS",
        "display": "Server", "table": "cmdb_ci_server",
        "domain": "Manage Technical Services",
        "in_scope": "Any active physical or virtual instance providing compute resources (regardless of OS).",
        "out_scope": ("Network appliances, storage arrays, decommissioned hardware, and end-user workstations "
                      "(laptops/desktops)."),
        "attrs": [
            ("Name", "name", "String (255)", "Y", "Free text", "ServiceNow Discovery", "—", "Y",
             "Ray Reuter (Class Mgr – Servers & Storage)", "CI identifier; IRE component."),
            ("Operational Status", "operational_status", "Choice", "Y", "Build, Operational, Non-Operational, Retired " + CONFIRM,
             "ServiceNow Discovery / Manual", "Operational", "Y",
             "Ray Reuter (Class Mgr – Servers & Storage)", "Lifecycle state. Confirm full allowed-value list against OOB choices."),
            ("IP Address", "ip_address", "String", "N", "Valid IPv4/IPv6", "ServiceNow Discovery (authoritative)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Servers & Storage)", "Primary network interface."),
            ("MAC Address", "mac_address", "String", "N", "Valid MAC format", "ServiceNow Discovery (authoritative)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Servers & Storage)", "Primary network interface; IRE component."),
            ("CPU Count", "cpu_count", "Integer", "N", "Numeric", "vCenter override (else Discovery)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Servers & Storage)", "vCenter overrides Discovery for CPU allocation only."),
            ("RAM (MB)", "ram", "Integer", "N", "Numeric (MB)", "vCenter override (else Discovery)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Servers & Storage)", "vCenter overrides Discovery for RAM allocation only."),
            ("Serial Number", "serial_number", "String", "Y", "Hardware serial", "ServiceNow Discovery", "—", "Y",
             "Ray Reuter (Class Mgr – Servers & Storage)", "IRE component (Serial + Type)."),
        ],
        "sources": "ServiceNow Discovery, VMware vCenter, Cloud Discovery (AWS/Azure).",
        "ire": "1. Serial Number + Type  |  2. MAC Address + Name  |  3. Name",
        "precedence": ("ServiceNow Discovery is authoritative for Hardware and Network attributes (ip_address, "
                       "mac_address). VMware vCenter integration overrides Discovery ONLY for CPU and RAM allocation "
                       "attributes."),
        "upstream": ["Used by :: Depends on → Application Service"],
        "downstream": ["Runs on :: Hosts → Virtual Machine Instance (if virtual) OR Hardware CI (if physical)"],
    },
    {
        "name": "Windows Server", "kind": "CHILD CLASS", "parent": "Server",
        "display": "Windows Server", "table": "cmdb_ci_win_server",
        "domain": "Manage Technical Services (inherited from Server)",
        "in_scope": "Compute instances running Microsoft Windows Server OS (e.g., 2016, 2019, 2022).",
        "out_scope": "Non-Windows compute instances.",
        "inherit_note": "Inherits all attributes, IRE, and relationships from parent class Server (cmdb_ci_server). Only attributes unique to or overridden by this child class are listed below.",
        "attrs": [
            ("OS Version", "os_version", "String", "N", "e.g., Windows Server 2019", "SCCM SGC (overrides Discovery)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Servers & Storage)", "SCCM overrides ServiceNow Discovery for OS attributes."),
            ("OS Domain", "os_domain", "String", "N", "AD domain", "SCCM SGC (overrides Discovery)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Servers & Storage)", "SCCM overrides Discovery."),
            ("OS Service Pack", "os_service_pack", "String", "N", "Free text", "SCCM SGC (overrides Discovery)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Servers & Storage)", "SCCM overrides Discovery."),
            ("ITSM Support Group", "support_group", "Reference (sys_user_group)", CONFIRM, "Wintel Infrastructure Ops", "Manual / CI Data Manager", "—", "Y",
             "CI Data Manager", "Governance attribute — NOT writable by automated tooling."),
            ("Change Approval Group", "change_control", "Reference (sys_user_group)", CONFIRM, "Wintel CAB", "Manual / CI Data Manager", "—", "Y",
             "CI Data Manager", "Governance attribute — NOT writable by automated tooling."),
        ],
        "sources": "Microsoft SCCM Service Graph Connector, ServiceNow Discovery (WMI).",
        "ire": "Inherited from Server.",
        "precedence": ("SCCM overrides ServiceNow Discovery for Operating System attributes (os_version, "
                       "os_service_pack, os_domain). Neither automated tool may overwrite Governance attributes "
                       "(support_group, change_control) — strictly manual / CI Data Manager."),
        "upstream": ["Inherited from Server"],
        "downstream": ["Inherited from Server"],
    },
    {
        "name": "Linux Server", "kind": "CHILD CLASS", "parent": "Server",
        "display": "Linux Server", "table": "cmdb_ci_linux_server",
        "domain": "Manage Technical Services (inherited from Server)",
        "in_scope": "Compute instances running a Linux kernel (e.g., RHEL, Ubuntu, CentOS, SUSE).",
        "out_scope": "Non-Linux compute instances.",
        "inherit_note": "Inherits all attributes, IRE, and relationships from parent class Server (cmdb_ci_server). Only attributes unique to or overridden by this child class are listed below.",
        "attrs": [
            ("OS Architecture", "os_architecture", "String", "N", "e.g., x86_64", "ServiceNow Discovery (SSH, authoritative)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Servers & Storage)", "Discovery via SSH is authoritative for OS/Hardware."),
            ("Kernel Release", "kernel_release", "String", "N", "Free text", "ServiceNow Discovery (SSH, authoritative)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Servers & Storage)", "Discovery via SSH is authoritative."),
            ("ITSM Support Group", "support_group", "Reference (sys_user_group)", CONFIRM, "Unix/Linux Administration Team", "Manual / CI Data Manager", "—", "Y",
             "CI Data Manager", "Governance attribute — manual / Data Manager control."),
            ("Change Approval Group", "change_control", "Reference (sys_user_group)", CONFIRM, "Linux CAB", "Manual / CI Data Manager", "—", "Y",
             "CI Data Manager", "Governance attribute — manual / Data Manager control."),
        ],
        "sources": "ServiceNow Discovery (SSH).",
        "ire": "Inherited from Server.",
        "precedence": ("ServiceNow Discovery via SSH is strictly authoritative for all OS and Hardware attributes. "
                       "Governance attributes (support_group, change_control) remain under manual / Data Manager control."),
        "upstream": ["Inherited from Server"],
        "downstream": ["Inherited from Server"],
    },
    {
        "name": "Database", "kind": "PARENT CLASS",
        "display": "Database Instance", "table": "cmdb_ci_database",
        "domain": "Manage Technical Services",
        "in_scope": "Logical database instances providing data storage and retrieval capabilities.",
        "out_scope": "The underlying host server OS, or the underlying physical storage / disks.",
        "attrs": [
            ("Name", "name", "String (255)", "Y", "Free text", "ServiceNow Discovery (App Probes)", "—", "Y",
             "Ray Reuter (Class Mgr – Database CIs)", "CI identifier."),
            ("Version", "version", "String", "N", "Engine version string", "ServiceNow Discovery (authoritative)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Database CIs)", "Active running version of the DB engine."),
            ("TCP Port", "tcp_port", "Integer", "N", "Numeric port", "ServiceNow Discovery (authoritative)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Database CIs)", "Listening port; IRE component."),
            ("Operational Status", "operational_status", "Choice", "Y", "Operational, Retired, … " + CONFIRM, "ServiceNow Discovery / Manual", "Operational", "Y",
             "Ray Reuter (Class Mgr – Database CIs)", "Lifecycle state."),
        ],
        "sources": "ServiceNow Discovery (Application Probes), Database Management integrations.",
        "ire": "Engine-dependent; usually Port + running process on the Host.",
        "precedence": "ServiceNow Discovery Application Probes are authoritative for active running version and port details.",
        "upstream": ["Used by :: Depends on → Application Service / Service Instance"],
        "downstream": ["Runs on :: Hosts → Server (Windows or Linux)"],
    },
    {
        "name": "Oracle Database", "kind": "CHILD CLASS", "parent": "Database",
        "display": "Oracle Instance", "table": "cmdb_ci_db_ora_instance",
        "domain": "Manage Technical Services (inherited from Database)",
        "in_scope": "Logical instances running the Oracle RDBMS engine.",
        "out_scope": "Other database engines.",
        "inherit_note": "Inherits all attributes, IRE, and relationships from parent class Database (cmdb_ci_database). Only attributes unique to this child class are listed below.",
        "attrs": [
            ("SID", "sid", "String", "Y", "Oracle SID", "ServiceNow Discovery (Oracle probes)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Database CIs)", "System identifier for the Oracle instance."),
            ("Oracle Edition", "edition", "String", "N", "e.g., Enterprise, Standard", "ServiceNow Discovery (Oracle probes)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Database CIs)", "Oracle edition."),
            ("Oracle Version", "version", "String", "N", "Version string", "ServiceNow Discovery (Oracle probes)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Database CIs)", "Oracle engine version."),
            ("Support Group", "support_group", "Reference (sys_user_group)", CONFIRM, "Oracle DBA Team", "Manual / CI Data Manager", "—", "Y",
             "CI Data Manager", "Governance attribute."),
        ],
        "sources": "ServiceNow Discovery via Oracle-specific probes / patterns.",
        "ire": "Inherited from Database.",
        "precedence": "Oracle DB-specific patterns override basic TCP-port discovery mapping for deeper instance metrics.",
        "upstream": ["Inherited from Database"],
        "downstream": ["Contains :: Contained By → Oracle Catalogs (cmdb_ci_db_ora_catalog)"],
    },
    {
        "name": "MSFT SQL Database", "kind": "CHILD CLASS", "parent": "Database",
        "display": "MSFT SQL Instance", "table": "cmdb_ci_db_mssql_instance",
        "domain": "Manage Technical Services (inherited from Database)",
        "in_scope": "Logical instances running Microsoft SQL Server.",
        "out_scope": "Other database engines.",
        "inherit_note": "Inherits all attributes, IRE, and relationships from parent class Database (cmdb_ci_database). Only attributes unique to this child class are listed below.",
        "attrs": [
            ("Instance Name", "instance_name", "String", "Y", "Free text", "ServiceNow Discovery / SCCM SGC", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Database CIs)", "SQL instance name."),
            ("Clustered", "clustered", "True/False", "N", "true / false", "ServiceNow Discovery (authoritative for clustering)", "false", CONFIRM,
             "Ray Reuter (Class Mgr – Database CIs)", "Whether the instance is clustered."),
            ("SQL Edition", "edition", "String", "N", "e.g., Enterprise, Standard", "SCCM SGC (authoritative for edition/license)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Database CIs)", "SCCM is authoritative for Edition and License parameters."),
            ("TCP Port", "tcp_port", "Integer", "N", "Numeric port", "ServiceNow Discovery (overrides SCCM)", "—", CONFIRM,
             "Ray Reuter (Class Mgr – Database CIs)", "Discovery overrides SCCM for active running Port and clustering metrics."),
            ("Support Group", "support_group", "Reference (sys_user_group)", CONFIRM, "SQL DBA Team", "Manual / CI Data Manager", "—", "Y",
             "CI Data Manager", "Governance attribute."),
        ],
        "sources": "ServiceNow Discovery (WMI/PowerShell), SCCM Service Graph Connector.",
        "ire": "Inherited from Database.",
        "precedence": ("SCCM Service Graph Connector is authoritative for SQL Server Edition and License parameters. "
                       "ServiceNow Discovery overrides SCCM for active running Port and clustering metrics."),
        "upstream": ["Inherited from Database"],
        "downstream": ["Contains :: Contained By → MSFT SQL Databases (cmdb_ci_db_mssql_database)"],
    },
]

# ---------------------------------------------------------------- Markdown

def md_escape(s):
    return s.replace("|", "\\|")

def render_md():
    L = []
    L.append("---")
    L.append("workspace: Work")
    L.append("tags: [cmdb-csdm, servicenow, governance, ccb, data-definitions]")
    L.append(f"updated: 2026-06-24")
    L.append("---")
    L.append("")
    L.append("# Configuration Management Plan")
    L.append("## Stage 1: Class & Attribute Data Definitions")
    L.append("")
    L.append("> **Audience:** CMDB Change Control Board (CCB). **Status:** DRAFT for CCB review.")
    L.append(">")
    L.append("> This document defines the Stage 1 Data Definitions for key CSDM and CMDB classes, including an "
             "attribute-level data dictionary, explicit Data Precedence / Reconciliation Rules, and CSDM relationship "
             "mappings establishing the source of truth for each attribute across discovery sources.")
    L.append(">")
    L.append("> **Cells marked `[CCB confirm]`** are PPL-specific governance decisions (mandatory rules, audit "
             "inclusion, certification ownership) that require CCB ratification — they are intentionally not assumed. "
             "**Certification Owner** is pre-populated from the CCB Class Manager roster where a class mapping exists.")
    L.append("")
    L.append("**Data dictionary columns:** " + " · ".join(COLS))
    L.append("")
    L.append("---")
    L.append("")

    for c in CLASSES:
        kind = c["kind"]
        header = f"## {kind}: {c['name']}"
        L.append(header)
        if c.get("parent"):
            L.append(f"*Child of parent class **{c['parent']}**.*")
        L.append("")

        # 1. Identity
        L.append("### 1. Class Identity & Scope")
        L.append("")
        L.append(f"- **Class Display Name:** {c['display']}")
        L.append(f"- **ServiceNow Table Name:** `{c['table']}`")
        L.append(f"- **CSDM Domain:** {c['domain']}")
        L.append(f"- **In Scope:** {c['in_scope']}")
        L.append(f"- **Out of Scope:** {c['out_scope']}")
        L.append("")

        # 2. Data dictionary
        L.append("### 2. Attribute-Level Data Dictionary")
        L.append("")
        if c.get("inherit_note"):
            L.append(f"> {c['inherit_note']}")
            L.append("")
        L.append("| " + " | ".join(COLS) + " |")
        L.append("|" + "|".join(["---"] * len(COLS)) + "|")
        for a in c["attrs"]:
            display, technical, dtype, mand, allowed, sot, default, audited, cert, desc = a
            row = [display, f"`{technical}`", dtype, mand, allowed, sot, default, audited, cert, DRAFT, desc]
            L.append("| " + " | ".join(md_escape(str(x)) for x in row) + " |")
        L.append("")

        # 3. Population
        L.append("### 3. Data Population & Identification")
        L.append("")
        L.append(f"- **Primary Discovery Source(s):** {c['sources']}")
        L.append(f"- **Identification Rules (IRE):** {c['ire']}")
        L.append(f"- **Data Precedence / Reconciliation Rules:** {c['precedence']}")
        L.append("")

        # 4. Relationships
        L.append("### 4. Relationship Architecture (CSDM Mapping)")
        L.append("")
        if c.get("rel_note"):
            L.append(f"- {c['rel_note']}")
        else:
            up = c.get("upstream") or ["None defined at Stage 1."]
            down = c.get("downstream") or ["None defined at Stage 1."]
            L.append("- **Upstream Relationships:**")
            for r in up:
                L.append(f"  - {r}")
            L.append("- **Downstream Relationships:**")
            for r in down:
                L.append(f"  - {r}")
        L.append("")
        L.append("---")
        L.append("")

    out = os.path.join(HERE, "configuration-management-plan-stage1.md")
    with open(out, "w", encoding="utf-8") as f:
        f.write("\n".join(L))
    return out

# ---------------------------------------------------------------- DOCX

def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=7, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def render_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(0.4))

    title = doc.add_heading("Configuration Management Plan", level=0)
    doc.add_heading("Stage 1: Class & Attribute Data Definitions", level=1)
    p = doc.add_paragraph()
    p.add_run("Audience: ").bold = True
    p.add_run("CMDB Change Control Board (CCB).    ")
    p.add_run("Status: ").bold = True
    p.add_run("DRAFT for CCB review.")
    doc.add_paragraph(
        "This document defines the Stage 1 Data Definitions for key CSDM and CMDB classes, including an "
        "attribute-level data dictionary, explicit Data Precedence / Reconciliation Rules, and CSDM relationship "
        "mappings establishing the source of truth for each attribute across discovery sources."
    )
    note = doc.add_paragraph()
    note.add_run("Cells marked [CCB confirm] ").bold = True
    note.add_run(
        "are PPL-specific governance decisions (mandatory rules, audit inclusion, certification ownership) requiring "
        "CCB ratification. Certification Owner is pre-populated from the CCB Class Manager roster where a mapping exists."
    )

    HEADER_FILL = "1F3864"   # dark blue
    KIND_FILL = "D9E1F2"     # light blue band

    for c in CLASSES:
        doc.add_page_break()
        h = doc.add_heading(f"{c['kind']}: {c['name']}", level=1)
        if c.get("parent"):
            sub = doc.add_paragraph()
            r = sub.add_run(f"Child of parent class {c['parent']}.")
            r.italic = True

        # 1. Identity
        doc.add_heading("1. Class Identity & Scope", level=2)
        idt = doc.add_table(rows=5, cols=2)
        idt.style = "Light Grid Accent 1"
        idrows = [
            ("Class Display Name", c["display"]),
            ("ServiceNow Table Name", c["table"]),
            ("CSDM Domain", c["domain"]),
            ("In Scope", c["in_scope"]),
            ("Out of Scope", c["out_scope"]),
        ]
        for i, (k, v) in enumerate(idrows):
            set_cell_text(idt.rows[i].cells[0], k, bold=True, size=9)
            set_cell_text(idt.rows[i].cells[1], v, size=9)
            idt.rows[i].cells[0].width = Inches(2.0)
            idt.rows[i].cells[1].width = Inches(8.0)

        # 2. Data dictionary
        doc.add_heading("2. Attribute-Level Data Dictionary", level=2)
        if c.get("inherit_note"):
            ip = doc.add_paragraph()
            ip.add_run(c["inherit_note"]).italic = True

        t = doc.add_table(rows=1, cols=len(COLS))
        t.style = "Table Grid"
        t.autofit = True
        for j, col in enumerate(COLS):
            set_cell_text(t.rows[0].cells[j], col, bold=True, size=7, color=RGBColor(0xFF, 0xFF, 0xFF))
            shade_cell(t.rows[0].cells[j], HEADER_FILL)
        for a in c["attrs"]:
            display, technical, dtype, mand, allowed, sot, default, audited, cert, desc = a
            vals = [display, technical, dtype, mand, allowed, sot, default, audited, cert, DRAFT, desc]
            cells = t.add_row().cells
            for j, v in enumerate(vals):
                set_cell_text(cells[j], v, size=7)

        # 3. Population
        doc.add_heading("3. Data Population & Identification", level=2)
        pt = doc.add_table(rows=3, cols=2)
        pt.style = "Light Grid Accent 1"
        prows = [
            ("Primary Discovery Source(s)", c["sources"]),
            ("Identification Rules (IRE)", c["ire"]),
            ("Data Precedence / Reconciliation Rules", c["precedence"]),
        ]
        for i, (k, v) in enumerate(prows):
            set_cell_text(pt.rows[i].cells[0], k, bold=True, size=9)
            set_cell_text(pt.rows[i].cells[1], v, size=9)
            pt.rows[i].cells[0].width = Inches(2.5)
            pt.rows[i].cells[1].width = Inches(7.5)

        # 4. Relationships
        doc.add_heading("4. Relationship Architecture (CSDM Mapping)", level=2)
        if c.get("rel_note"):
            doc.add_paragraph(c["rel_note"])
        else:
            up = c.get("upstream") or ["None defined at Stage 1."]
            down = c.get("downstream") or ["None defined at Stage 1."]
            pu = doc.add_paragraph()
            pu.add_run("Upstream Relationships:").bold = True
            for r in up:
                doc.add_paragraph(r, style="List Bullet")
            pd = doc.add_paragraph()
            pd.add_run("Downstream Relationships:").bold = True
            for r in down:
                doc.add_paragraph(r, style="List Bullet")

    out = os.path.join(HERE, "Configuration-Management-Plan-Stage1.docx")
    doc.save(out)
    return out

if __name__ == "__main__":
    md = render_md()
    print("Wrote:", md)
    dx = render_docx()
    print("Wrote:", dx)
    print(f"Classes: {len(CLASSES)} | Columns: {len(COLS)}")
