"""
CMDB Computer & Server Record Cleanup — Root Cause, Impact, and Response.
Word (.docx) output, mirroring cmdb-cleanup-root-cause-response-2026-07-23.md
(the markdown remains the single source of truth — keep them in sync).

Run:     python build_cleanup_response_doc.py
Output:  CMDB-Cleanup-Root-Cause-Response-2026-07-23.docx
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

EDITION = "2026-07-23"
OUTFILE = "CMDB-Cleanup-Root-Cause-Response-" + EDITION + ".docx"

DARK   = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x6D, 0xB4)
MUTED  = RGBColor(0x6B, 0x72, 0x80)
RED    = RGBColor(0xC0, 0x39, 0x2B)

# --------------------------------------------------------------------------- CONTENT
NOTE = ('DRAFT, for review before sending. Audience: Sonika first, then broader key stakeholders. '
        'Purpose: a root-cause and remediation summary for leadership, and a response to the concerns '
        'raised in Sonika’s email thread ("RE: Issues...SNOW CMDB").')

OPEN_ITEMS = [
    'Record count: currently stated as "all qualifying records." Insert the actual count from the query result if you want it shown.',
    "Remediation (Concerns #6/#8): now framed as a governed CMDB Health and retirement lifecycle policy "
    "(see Governance & Go-Forward Controls). Requirements still need to be defined with stakeholders and "
    "approved by the CCB before development.",
]

EXEC = [
    "We ran a one-time cleanup to retire Computer and Server records that were loaded into the CMDB from legacy "
    "sources as an initial starting point, before ServiceNow Discovery scope was finalized. Concerns were then raised "
    "that active operational assets may have been caught in that cleanup, with risk to data integrity, operations, "
    "governance, and compliance.",
    "The main cause is a gap between what the CMDB actually covers today and how broadly its coverage is understood. "
    "Discovery is intentionally limited to specific subnets. Many subnets, including those hosting OT, SCADA, and "
    "NERC-related devices, are out of scope and are not discovered or maintained in the CMDB. The cleanup removed "
    "unvalidated pre-scope import records, not live discovered assets.",
    "Every retirement is **fully reversible**. Retired records are not deleted, so any record later confirmed as a "
    "legitimate in-scope asset can be restored immediately with no data loss. Going forward, a governed CMDB Health "
    "and lifecycle policy, defined with stakeholder input and approved by the CCB, will gate all future retirements "
    "(see Governance & Go-Forward Controls). The current focus is a stable, in-scope CMDB baseline.",
]

WHAT_OCCURRED = [
    "A single, one-time query identified and retired the qualifying records. This was **not an automated or recurring "
    "process** writing to the CMDB on a schedule.",
    "The criteria: records not already Retired whose Discovery Source was a legacy or manual origin (ImportSet, "
    "iTeam, Manual Entry, CherWell KY, CherWell PA, or empty).",
    "All records meeting the criteria were retired, across both the Computer and Server classes.",
    "The action was **reviewed with key stakeholders in advance**. The CMDB team did not do this unilaterally.",
]

ROOT_CAUSE = [
    "The initial CMDB loads were bulk collections of records from existing sources, imported as a starting point "
    "before Discovery scope was finalized. Because scope was not yet defined, these records were **never validated "
    "for scope** when they entered the CMDB. They are not maintained by any authoritative discovery source, and many "
    "sit outside the subnets the CMDB is scoped to manage.",
    "The concern that active operational assets were retired comes largely from that same scope gap. The CMDB covers "
    "a narrower slice of the environment today than most people assume.",
]

SCOPE = [
    "ServiceNow Discovery is scoped to specific, prescribed subnets.",
    "Multiple subnets are **intentionally out of scope**, including subnets that host OT, SCADA, and NERC-related "
    "devices. By design, those assets are not discovered and not maintained in the CMDB.",
    "Retirement logic based on the absence of active discovery does not act on those out-of-scope operational assets, "
    "because they are not in the discovered inventory in the first place.",
    "**NERC CIP is currently out of scope** for this CMDB effort. A separate, parallel initiative is evaluating "
    "NERC-specific solutions.",
]

CONCERNS = [
    (1, "Active CIs automatically retired",
     "Done by a **one-time script**, not an automated or recurring rule. Reviewed with key stakeholders beforehand, "
     "not unilaterally."),
    (2, "Potential widespread / undiscovered impact",
     "Retired records were unvalidated pre-scope imports, not authoritatively discovered assets. No affected record "
     "has been identified as a legitimate in-scope asset. Any that turns out to be one is **restored immediately**, "
     "since retirement is reversible."),
    (3, "Communication / stakeholder awareness",
     "The cleanup was reviewed with key stakeholders. Communication is deliberately staged: key stakeholders now to "
     "stabilize the baseline, then operational teams as the baseline is confirmed."),
    (4, "Operational risk (monitoring alerts)",
     "Where a retired record was still referenced downstream, discrepancies could surface. Restoring the record "
     "resolves them, and **over 40 records found this way have already been restored**. These cases also help confirm "
     "which records belong in the validated baseline."),
    (5, "OT / SCADA / CIP assets are different",
     "Agreed, and this is central. Those assets sit on **out-of-scope subnets**, are not discovered, and are not "
     "maintained in the CMDB. The cleanup targeted pre-scope import records, not live OT assets. OT and CIP are "
     "handled by a parallel effort."),
    (6, "Insufficient retirement controls",
     "Agreed. We are establishing a governed **CMDB Health and lifecycle policy**: staleness-based candidate "
     "identification, automated stakeholder notification, an explicit approval step, and advance notice before "
     "implementation (see Governance & Go-Forward Controls). Requirements will be defined with stakeholder input and "
     "**approved by the CCB** before any development."),
    (7, "Compliance / NERC CIP",
     "**NERC CIP is out of scope today**, and the CMDB is not its system of record. A parallel initiative owns "
     "NERC-specific solutions and will define governance for any future regulated data."),
    (8, "Root cause & remediation",
     "Root cause: unvalidated pre-scope legacy imports, plus a CMDB scope narrower than people assumed. Remediation: "
     "the governed CMDB Health retirement policy (see Governance & Go-Forward Controls). In the meantime, **full "
     "reversibility** (40+ already restored) and staged communication protect the move toward a validated baseline."),
]

SAFEGUARDS = [
    "Retired does not mean deleted. Every retired record is **fully recoverable**, and any legitimate in-scope asset "
    "can be restored on request. We have already done this in practice: **over 40 records restored to date**.",
    "The cleanup was a one-time baseline action, not a standing automated rule.",
    "Future retirements will run through the governed CMDB Health and lifecycle policy below, not ad-hoc scripts.",
]

GOV_INTRO = ("So this stays a one-time event and not a recurring risk, we are establishing a governed **CMDB Health "
             "and lifecycle policy** for identifying and retiring CIs. It puts every future retirement through a "
             "controlled, approval-gated workflow that leaves an audit trail:")

GOV_STEPS = [
    "**Candidate identification (CMDB Health).** Retirement candidates come from defined CMDB Health criteria, mainly "
    "staleness (CIs not updated or confirmed by an authoritative source within a set threshold) and related "
    "data-quality checks. The criteria will explicitly account for legitimately non-discoverable in-scope assets such "
    "as OT and SCADA, so that lack of discovery alone never triggers retirement.",
    "**Automated stakeholder notification.** When a CI meets the criteria, an automated workflow notifies the "
    "affected stakeholders and owners. No silent state changes.",
    "**Explicit approval step.** Retirement does not proceed without explicit approval from the responsible "
    "stakeholders and SMEs, keeping them in the lifecycle decision.",
    "**Advance notice before implementation.** Once approved, advance notice goes out before the change is made, "
    "giving downstream and operational teams time to complete their own decommissioning steps and avoid monitoring "
    "surprises.",
]

GOV_AUTHORITY = ("**Governance authority:** These policy requirements will be defined with stakeholder input and "
                 "formally approved by the Change Control Board (CCB) before any development or implementation.")
GOV_COMPLIANCE = ("**Process compliance:** The policy will comply with all existing change, configuration management, "
                  "and governance processes.")
GOV_LIFECYCLE = ("Lifecycle at a glance:  Identify (CMDB Health criteria)  →  Notify (automated)  →  "
                 "Approve (explicit)  →  Advance notice  →  Implement.")

NEXT_STEPS = [
    "Define the CMDB Health and retirement lifecycle policy with stakeholder input (staleness criteria, automated "
    "notification, explicit approval, advance notice) and **submit it to the CCB for approval before development**.",
    "Keep restoring any record identified as a legitimate in-scope asset (40+ done so far).",
    "Keep stabilizing the in-scope CMDB baseline, and widen communication as it firms up.",
]


# --------------------------------------------------------------------------- HELPERS
_TOKEN = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def add_rich(paragraph, text, size=None, color=None):
    """Render **bold** and *italic* markdown inline into runs."""
    for tok in _TOKEN.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        else:
            r = paragraph.add_run(tok)
        if size:
            r.font.size = size
        if color:
            r.font.color.rgb = color


def style_base(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(10.5)


def heading(doc, text, size, color=DARK, space_before=10, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(space_after)
    return p


def para(doc, text, space_after=6):
    p = doc.add_paragraph()
    add_rich(p, text)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    add_rich(p, text)
    p.paragraph_format.space_after = Pt(2)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    add_rich(p, text)
    p.paragraph_format.space_after = Pt(3)
    return p


def concerns_table(doc, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=3)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(0.35), Inches(2.05), Inches(4.20)]
    hdr = t.rows[0].cells
    for j, h in enumerate(["#", "Concern", "Response"]):
        hdr[j].text = ""
        r = hdr[j].paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = DARK
    for i, (num, concern, resp) in enumerate(rows, start=1):
        cells = t.rows[i].cells
        cells[0].text = ""
        rn = cells[0].paragraphs[0].add_run(str(num)); rn.font.size = Pt(9); rn.bold = True
        cells[1].text = ""; add_rich(cells[1].paragraphs[0], concern, Pt(9))
        cells[2].text = ""; add_rich(cells[2].paragraphs[0], resp, Pt(9))
    for i in range(len(rows) + 1):
        for j, w in enumerate(widths):
            t.rows[i].cells[j].width = w
    return t


# --------------------------------------------------------------------------- BUILD
def build():
    doc = Document()
    style_base(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)

    # Title block
    title = doc.add_paragraph()
    r = title.add_run("ServiceNow CMDB: Computer & Server Record Cleanup"); r.bold = True
    r.font.size = Pt(19); r.font.color.rgb = DARK
    sub = doc.add_paragraph()
    r = sub.add_run("Root Cause, Impact, and Response to Concerns Raised")
    r.font.size = Pt(13); r.font.color.rgb = ACCENT
    meta = doc.add_paragraph()
    r = meta.add_run("PPL CMDB-CSDM · DRAFT · " + EDITION)
    r.font.size = Pt(9); r.font.color.rgb = MUTED; r.italic = True

    # Draft note
    dn = doc.add_paragraph()
    r = dn.add_run("DRAFT: For Review Before Sending"); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RED
    note = doc.add_paragraph()
    add_rich(note, NOTE); note.runs[0].italic = True
    for r_ in note.runs:
        r_.font.size = Pt(9); r_.font.color.rgb = MUTED; r_.italic = True
    oi = doc.add_paragraph()
    r = oi.add_run("Open items before sending:"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = DARK
    for item in OPEN_ITEMS:
        p = doc.add_paragraph(style="List Bullet")
        add_rich(p, item)
        for r_ in p.runs:
            r_.font.size = Pt(9); r_.font.color.rgb = MUTED
        p.paragraph_format.space_after = Pt(2)

    # Executive Summary
    heading(doc, "Executive Summary", 14)
    for t in EXEC:
        para(doc, t)

    # What Occurred
    heading(doc, "What Occurred", 14)
    for t in WHAT_OCCURRED:
        bullet(doc, t)

    # Root Cause
    heading(doc, "Root Cause", 14)
    for t in ROOT_CAUSE:
        para(doc, t)

    # Scope Clarification
    heading(doc, "Scope Clarification", 14)
    for t in SCOPE:
        bullet(doc, t)

    # Response to Concerns
    heading(doc, "Response to the Concerns Raised", 14)
    concerns_table(doc, CONCERNS)

    # Safeguards
    heading(doc, "Safeguards & Reversibility", 14)
    for t in SAFEGUARDS:
        bullet(doc, t)

    # Governance
    heading(doc, "Governance & Go-Forward Controls", 14)
    para(doc, GOV_INTRO)
    for t in GOV_STEPS:
        numbered(doc, t)
    para(doc, GOV_AUTHORITY, space_after=3)
    para(doc, GOV_COMPLIANCE, space_after=6)
    lc = doc.add_paragraph()
    r = lc.add_run(GOV_LIFECYCLE); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = ACCENT
    lc.paragraph_format.space_before = Pt(2); lc.paragraph_format.space_after = Pt(6)

    # Next Steps
    heading(doc, "Next Steps", 14)
    for t in NEXT_STEPS:
        numbered(doc, t)

    return doc


def save_safe(doc, path):
    try:
        doc.save(path); print("Saved", os.path.basename(path))
    except PermissionError:
        base, ext = os.path.splitext(path)
        alt = base + "-new" + ext
        doc.save(alt); print("Target locked — saved to", os.path.basename(alt))


here = os.path.dirname(os.path.abspath(__file__))
save_safe(build(), os.path.join(here, OUTFILE))
