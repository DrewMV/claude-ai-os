"""
CO5 Alignment & Deliverable Health — leadership brief (Word document).

Single source of truth for the .docx. Edit CONTENT below, then run:
    python build_co5_brief_doc.py
Output: 2026-06-23-co5-alignment-brief.docx

Mirrors the humanized markdown brief in this folder. Theme colors match the
status/scope decks so the artifacts look like a set.
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- theme (mirrors build_status_deck.py) -------------------------------
DARK  = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT= RGBColor(0x2E, 0x6D, 0xB4)
TEXT  = RGBColor(0x2D, 0x2D, 0x2D)
MUTED = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = "2E8B57"
AMBER = "E08A00"
RED   = "C0392B"
GREY  = "9AA3AD"
HDR   = "1F3A5F"
ZEBRA = "F4F6F9"
FONT  = "Calibri"

RAG = {"G": GREEN, "A": AMBER, "R": RED, "NA": GREY}
RAG_LABEL = {"G": "G", "A": "A", "R": "R", "NA": "n/a"}

# --- helpers ------------------------------------------------------------
def add_md_runs(p, text, size=11, color=TEXT, italic=False):
    """Render a string with **bold** markers into runs."""
    for i, part in enumerate(re.split(r'(\*\*)', text)):
        if part == '**' or part == '':
            continue
    # rebuild properly toggling bold
    p.clear() if False else None
    bold = False
    for part in re.split(r'(\*\*)', text):
        if part == '**':
            bold = not bold
            continue
        if part == '':
            continue
        r = p.add_run(part)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = FONT

def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexfill)
    tcPr.append(shd)

def set_cell(cell, text, size=10, color=TEXT, bold=False, fill=None,
             align=WD_ALIGN_PARAGRAPH.LEFT):
    if fill:
        shade(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    add_md_runs(p, text, size=size, color=color)
    # bold override (whole cell) when requested and no ** markers
    if bold and '**' not in text:
        for r in p.runs:
            r.bold = True

def set_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for c, w in zip(row.cells, widths):
            c.width = Inches(w)
    for col, w in zip(table.columns, widths):
        col.width = Inches(w)

def header_row(table, labels):
    for j, lab in enumerate(labels):
        c = table.rows[0].cells[j]
        set_cell(c, lab, size=10, color=WHITE, fill=HDR)
        for r in c.paragraphs[0].runs:
            r.bold = True

def zebra(table, start=1):
    for i, row in enumerate(table.rows):
        if i < start:
            continue
        if (i - start) % 2 == 1:
            for c in row.cells:
                if not c._tc.findall(qn('w:tcPr') + '/' + qn('w:shd')):
                    shade(c, ZEBRA)

def heading(doc, text, size=15, color=DARK, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = FONT
    return p

def body(doc, text, size=11, italic=False, color=TEXT, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    add_md_runs(p, text, size=size, color=color, italic=italic)
    return p

def bullet(doc, text, size=11, num=False):
    style = 'List Number' if num else 'List Bullet'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    add_md_runs(p, text, size=size)
    return p

def callout(doc, text):
    """Light-bordered info paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    add_md_runs(p, text, size=10.5, color=RGBColor(0x33, 0x44, 0x55), italic=True)
    return p

# --- document -----------------------------------------------------------
doc = Document()
doc.styles['Normal'].font.name = FONT
doc.styles['Normal'].font.size = Pt(11)
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.7)
    s.top_margin = s.bottom_margin = Inches(0.7)

# Title
t = doc.add_paragraph()
r = t.add_run("PI-2 Status Update — CO5 Alignment & Deliverable Health")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = DARK; r.font.name = FONT
m = doc.add_paragraph()
add_md_runs(m, "**To:** CMDB-CSDM Leadership    **From:** Manuel Vazquez (SM / PO Support)", size=10.5, color=MUTED)
m2 = doc.add_paragraph()
add_md_runs(m2, "**As of:** Jun 23, 2026 — Iteration 2.3 close · ADO pull 6/23", size=10.5, color=MUTED)

# Bottom Line
heading(doc, "Bottom Line")
for b in [
    "**About a third of the delivery effort we've spent through Sprint 2.3 went to work that isn't in CO5.** That holds up whichever way I measure it: ~**30% by story/spike count**, ~**33% by points**. Airlift by itself is two people for the whole sprint on work that needs a change order.",
    "**Point tracking this PI is uneven** — only ~74% of user stories are pointed, and the tasks that make up most of Service Mapping aren't pointed at all. So I'm leading with a count of stories and spikes and using points only as a cross-check. The headline doesn't rest on the points either way.",
    "**All three CO5 deliverables are due June 30, a week out.** Four hard gaps are still open, and the **$533,775 holdback** depends on accepting them.",
    "The added work, Airlift and NowAssist, has been absorbed into the team's capacity **with no change order**. CO5 Article 2 says scope changes have to be in writing.",
]:
    bullet(doc, b)

# Section 1
heading(doc, "1. PI-2 effort out of scope for CO5 — through Sprint 2.3")
callout(doc, "Why I'm counting items, not just points: point tracking this PI is uneven. Only 29 of 39 user stories (74%) carry points, and tasks — the biggest category of work, and nearly all of Service Mapping — don't carry points at all by design. Measure by points alone and you miss most of what the team actually worked. So the main view below counts stories and spikes; points are there to back it up. This is Iterations 2.1–2.3 only, since 2.4–2.6 haven't run yet.")
body(doc, "**Both ways of counting land in the same place: about 30–33% out of scope.**")

# Table A — both measures
tA = doc.add_table(rows=3, cols=2); tA.style = 'Table Grid'
header_row(tA, ["Measure (through Sprint 2.3)", "Out of scope for CO5"])
set_cell(tA.rows[1].cells[0], "By story / spike count")
set_cell(tA.rows[1].cells[1], "~15 of ~51 ≈ **30%**")
set_cell(tA.rows[2].cells[0], "By story points")
set_cell(tA.rows[2].cells[1], "24 of 72 ≈ **33%**")
set_widths(tA, [3.6, 3.5]); zebra(tA)

heading(doc, "Primary view — by story / spike count", size=12, space_before=12)
tB = doc.add_table(rows=7, cols=3); tB.style = 'Table Grid'
header_row(tB, ["Category", "Stories + Spikes", "Share"])
rowsB = [
    ("In-scope — CO5 deliverables (18 stories + 6 spikes)", "24", ""),
    ("In-scope — Service Mapping (8 + 1)", "9", ""),
    ("**In-scope subtotal**", "**33**", "**~65%**"),
    ("Out — **Airlift** (4)¹ · **NowAssist** (5) · **Qualys build** (2) · **Network Gear** (2) · Sailpoint / Moveworks (2)", "**15**", "**~30%**"),
    ("Borderline — Operational Monitoring (2 + 1)", "3", "~6%"),
    ("**Total**", "**~51**", ""),
]
for i, (a, b, c) in enumerate(rowsB, start=1):
    set_cell(tB.rows[i].cells[0], a)
    set_cell(tB.rows[i].cells[1], b, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(tB.rows[i].cells[2], c, align=WD_ALIGN_PARAGRAPH.CENTER)
set_widths(tB, [4.3, 1.5, 1.3]); zebra(tB)
body(doc, "¹ Airlift's 4 stories sit past the export slice but are counted per the flagged 2-resource Sprint-2.3 effort.", size=9, color=MUTED)

heading(doc, "Corroborating view — by story points", size=12, space_before=12)
tC = doc.add_table(rows=5, cols=4); tC.style = 'Table Grid'
header_row(tC, ["Out-of-scope stream", "Pts", "Items", "Status"])
rowsC = [
    ("**Airlift / VMware→Azure** (P0 · Feat 1420613)", "**16**", "1418610 / 18 / 21 + 1416384 (unpointed; 16 = 2 resources × S2.3)", "3 Validation · 1 Active"),
    ("**NowAssist AI** (P5 · Feat 1436574)", "**6**", "1436576, 1436579, 1436592, 1436593 + 1470837 health spike", "Mostly Resolved / Validation"),
    ("**Qualys build** (P4 · beyond CO5 requirements)", "**2**", "1428703, 1428704", "BLOCKED — vendor plugin (Issue 1465952)"),
    ("**Out-of-scope subtotal**", "**24**", "vs **48** in-scope (CO5 35 + Service Mapping 13)", ""),
]
for i, (a, b, c, d) in enumerate(rowsC, start=1):
    set_cell(tC.rows[i].cells[0], a)
    set_cell(tC.rows[i].cells[1], b, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(tC.rows[i].cells[2], c, size=9)
    set_cell(tC.rows[i].cells[3], d, size=9)
set_widths(tC, [2.2, 0.5, 2.6, 1.8]); zebra(tC)
body(doc, "Both numbers are floors. Network Gear (CO6), NERC-CIP, and Upgrade Analysis are out of scope too, but they're mostly unpointed, so they barely show up in the points column. The in-scope side is undercounted as well: a few pointed 2.3 stories (SCCM Server precedence, 1454371, 1387236, ~7 pts) landed after my export slice. Add those and in-scope rises, which pushes out-of-scope a bit under 30%.", size=10, italic=True, color=RGBColor(0x33,0x44,0x55))

# Section 2
heading(doc, "2. CO5 deliverable status — all due June 30, 2026")
body(doc, "Accepting these three deliverables is what releases the **$533,775 holdback**.")

co5 = [
    ("1. Governance", "1.1 Data Dictionary — Servers", "Jun 30", "A", "Feat 1354797 Ready; 1421790 Validation; 1454371 Active"),
    ("", "1.1 — Computers", "Jun 30", "A", "1313400 Resolved; 1399547 / 1455858 Validation"),
    ("", "1.1 — Business Apps", "Jun 30", "A", "1475584/585/582, 1478286, 1474892 — New / Refinement (early)"),
    ("", "1.1 — Databases [G1]", "Jun 30", "R", "HARD GAP — no dedicated DB data-dictionary story"),
    ("", "1.2 Data Certification", "Jun 30", "G", "Dashboard 1402727 UAT signed (6/10), PROD 6/23; 1435307 Resolved"),
    ("", "1.3 ESS-02 Alignment", "Jun 30", "A", "Spike 1420244 Active — analysis only; parent 1406668 Removed"),
    ("", "1.4 SOX BA review (SOX only, not NERC/CIP)", "Jun 30", "A", "Issue 1438967 Closed/approved; 1455827 in Iter 2.5 (post-deadline)"),
    ("", "1.5 Monthly CCB + future backlog", "Jun 30", "R", "Cadence running (CCB 6/16) but 1406672/683/687 under Removed feature 1406668"),
    ("2. Automated Data Ingestion", "2.1 Computers — SCCM/Discovery precedence", "Jun 30", "G", "1348712/716/717 Resolved; 1348715 Validation"),
    ("", "2.1 Computers — 90% coverage [G2]", "Jun 30", "R", "HARD GAP — no coverage-measurement story"),
    ("", "2.1 Computers — bulk Lifecycle", "Jun 30", "G", "1402790 Closed"),
    ("", "2.2 Servers — SCCM/Discovery precedence", "Jun 30", "G", "1403759/760/762/763 Validation"),
    ("", "2.2 Servers — credentials enabling discovery", "Jun 30", "A", "1444864 Validation; child tasks still Active"),
    ("", "2.2 Servers — 90% non-NERC-CIP coverage [G3]", "Jun 30", "R", "HARD GAP — no coverage-measurement story"),
    ("", "2.2 / 2.3 Servers & DB SOX indicators", "Jun 30", "NA", "Manual per CO5 — excluded from automation"),
    ("", "2.3 Databases — enhanced MS-SQL/Oracle Discovery", "Jun 30", "A", "Task-level only (creds under 1444864); no standalone story"),
    ("3. Other Enhancement", "Evaluate 1 integration (Qualys)", "Jun 30", "A", "Requirements 1234585 Ready; build BLOCKED (1428703/704, Issue 1465952)"),
]
tD = doc.add_table(rows=len(co5)+1, cols=5); tD.style = 'Table Grid'
header_row(tD, ["Deliverable", "Sub-item", "Due", "RAG", "Status / evidence"])
for i, (deliv, sub, due, rag, status) in enumerate(co5, start=1):
    set_cell(tD.rows[i].cells[0], deliv, size=10, bold=True, color=DARK)
    set_cell(tD.rows[i].cells[1], sub, size=9.5)
    set_cell(tD.rows[i].cells[2], due, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    c = tD.rows[i].cells[3]
    set_cell(c, RAG_LABEL[rag], size=10, color=WHITE, fill=RAG[rag], align=WD_ALIGN_PARAGRAPH.CENTER)
    for rr in c.paragraphs[0].runs:
        rr.bold = True
    set_cell(tD.rows[i].cells[4], status, size=9.5)
set_widths(tD, [1.25, 1.95, 0.55, 0.5, 2.85])
# merge deliverable groups (col 0)
tD.cell(1, 0).merge(tD.cell(8, 0))   # Governance rows 1-8
tD.cell(9, 0).merge(tD.cell(16, 0))  # Auto Ingestion rows 9-16
for grp_top, label in [(1, "1. Governance"), (9, "2. Automated Data Ingestion")]:
    cell = tD.cell(grp_top, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].clear()
    set_cell(cell, "**" + label + "**", size=10, color=DARK)

body(doc, "**Scorecard:** 4 Green · 8 Amber · 4 Red · 1 n/a. The four reds are where acceptance is at risk: the **Databases data dictionary (G1)**, **90% coverage measurement for Computers and Servers (G2/G3)**, and **CCB evidence sitting under a feature marked Removed (1.5)**.", space_after=8)
body(doc, "Also in scope, but not a holdback-gated CO5 deliverable: Service Mapping. It's contracted under a prior CO and shows up in CO5 only in the PO role description on p.4, not the deliverables table. (Amber) Maps are moving (~13 pts pointed through 2.3), but about 36 map tasks are stranded in 2.1/2.2 carryover and Waves 20–21 are empty.", size=10, italic=True, color=RGBColor(0x33,0x44,0x55))

# Asks
heading(doc, "What we need from leadership")
for n in [
    "**A change-order decision on the added work.** Airlift (P0) and NowAssist are ~22 of the 24 out-of-scope points already spent through Sprint 2.3. CO5 Article 2 requires a written change order for scope changes, and right now this is being absorbed into a back half (Sprints 4–5) that's already at 100% capacity.",
    "**Sign CO6 before July 1.** CO6 re-baselines the at-risk CO5 items — Databases, 90% coverage, Qualys — to Jul 31 / Oct 30. If it isn't signed, the June 30 deadline and the $533,775 holdback stand with no relief.",
    "**Close or re-home the three hard gaps:** G1 (DB data dictionary) and G2 / G3 (coverage measurement). Either deliver them by 6/30 or move them into CO6.",
    "**Fix CCB feature 1406668** — un-Remove it or re-parent the stories — so the Deliverable 1.5 evidence isn't filed under a dead feature.",
    "**A ruling on SOX story 1455827.** It's in Iteration 2.5 right now, which falls after the 6/30 deadline.",
]:
    bullet(doc, n, num=True)

# Caveats
heading(doc, "Data provenance & caveats", size=12)
for cav in [
    "**Basis:** the 6/23 ADO export, Iterations 2.1–2.3 only (2.4–2.6 haven't run). Story/spike count is the primary measure; points corroborate.",
    "**Inconsistent point tracking (the key limitation):** 29 of 39 user stories (74%) are pointed, spikes are tracked well, and tasks and features carry no points by design. Tasks are the largest category of work and nearly all of Service Mapping. That's why count, not points, is the primary basis.",
    "**Airlift** isn't pointed in ADO. Its 16 pts / 4 stories reflect the 2-resource Sprint-2.3 effort, not ADO values.",
    "**These are floors, not ceilings.** Out-of-scope leaves out the largely-unpointed Network Gear, NERC-CIP, and Upgrade Analysis; in-scope leaves out ~7 pts of 2.3 stories past the export slice. The real out-of-scope share is probably a little under 30%.",
    "**Operational Monitoring** (3 items / 3 pts) is flagged pending a ruling. If it's ruled out of scope, out-of-scope rises to ~35% by count / 36% by points.",
    "**Service Mapping** counts as in-scope committed work per your direction, but it's not one of CO5's three holdback-gated deliverables.",
]:
    bullet(doc, cav, size=10)

# footer
sec = doc.sections[0]
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
fr = fp.add_run("PPL CMDB-CSDM  |  PI-2  |  CO5 Alignment Brief — Jun 23, 2026")
fr.font.size = Pt(8); fr.font.color.rgb = MUTED; fr.font.name = FONT


def save_safe(path):
    try:
        doc.save(path); print("Saved", path)
    except PermissionError:
        base, ext = os.path.splitext(path)
        alt = base + "-new" + ext
        doc.save(alt); print("Target locked — saved to", alt)


out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "2026-06-23-co5-alignment-brief.docx")
save_safe(out)
